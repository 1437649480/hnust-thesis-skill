# -*- coding: utf-8 -*-
"""
湖南科技大学本科毕业论文格式修复脚本
覆盖规范第3.3节、第4.1~4.2节所有可自动修复项。
用法: python fix_thesis.py <docx文件路径>
"""

import sys
import os
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── 常量定义 ───

# 4.1.1 页面设置: A4纸, 上30mm 下25mm 左25mm 右25mm, 页眉20mm 页脚15mm
PAGE_SPEC = {
    'width': Cm(21),        # A4宽度 210mm
    'height': Cm(29.7),     # A4高度 297mm
    'top': Cm(3.0),         # 上页边距 30mm
    'bottom': Cm(2.5),      # 下页边距 25mm
    'left': Cm(2.5),        # 左页边距 25mm
    'right': Cm(2.5),       # 右页边距 25mm
    'header': Cm(2.0),      # 页眉距边界 20mm
    'footer': Cm(1.5),      # 页脚距边界 15mm
}

# 4.1.2 页眉: 内容"湖南科技大学本科生毕业设计（论文）"，五号宋体(10.5pt)居中
HEADER_SPEC = {
    'text': '湖南科技大学本科生毕业设计（论文）',
    'font': '宋体',
    'size': Pt(10.5),       # 五号
    'align': WD_ALIGN_PARAGRAPH.CENTER,
}

# 4.2 段落格式规范
FONT_SPEC = {
    # 章标题: 第X章 ...  宋体加粗小二号(18pt), 居中, 段前段后各6pt
    'h1': {
        'font': '宋体', 'size': Pt(18), 'bold': True,
        'align': WD_ALIGN_PARAGRAPH.CENTER,
        'before': Pt(6), 'after': Pt(6),
    },
    # 节标题: X.Y ...  宋体加粗四号(14pt), 顶格, 段前段后各6pt
    'h2': {
        'font': '宋体', 'size': Pt(14), 'bold': True,
        'align': WD_ALIGN_PARAGRAPH.LEFT,
        'before': Pt(6), 'after': Pt(6),
    },
    # 条标题: X.Y.Z ...  宋体加粗小四号(12pt), 顶格, 段前段后各6pt
    'h3': {
        'font': '宋体', 'size': Pt(12), 'bold': True,
        'align': WD_ALIGN_PARAGRAPH.LEFT,
        'before': Pt(6), 'after': Pt(6),
    },
    # 款/项标题: (1) (一) 等  宋体加粗小四号(12pt), 段前段后不空, 1.25倍行距, 首行缩进2字符
    'h4': {
        'font': '宋体', 'size': Pt(12), 'bold': True,
        'align': WD_ALIGN_PARAGRAPH.LEFT,
        'before': Pt(0), 'after': Pt(0),
        'line_spacing': 1.25, 'first_line_indent': Pt(24),
    },
    # 正文: 宋体小四号(12pt), 段前段后不空, 1.25倍行距, 首行缩进2字符(24pt)
    'body': {
        'font': '宋体', 'size': Pt(12), 'bold': False,
        'align': WD_ALIGN_PARAGRAPH.LEFT,
        'before': Pt(0), 'after': Pt(0),
        'line_spacing': 1.25, 'first_line_indent': Pt(24),
    },
    # 参考文献标题: 宋体四号加粗(14pt), 居中
    'ref_title': {
        'font': '宋体', 'size': Pt(14), 'bold': True,
        'align': WD_ALIGN_PARAGRAPH.CENTER,
    },
    # 参考文献条目: 五号宋体(10.5pt), 顶格
    'ref_entry': {
        'font': '宋体', 'size': Pt(10.5), 'bold': False,
        'align': WD_ALIGN_PARAGRAPH.LEFT,
    },
    # 摘要标题: 宋体加粗小二号(18pt), 居中
    'abstract_title': {
        'font': '宋体', 'size': Pt(18), 'bold': True,
        'align': WD_ALIGN_PARAGRAPH.CENTER,
    },
    # ABSTRACT标题: Times New Roman加粗小二号(18pt), 居中
    'en_abstract_title': {
        'font': 'Times New Roman', 'size': Pt(18), 'bold': True,
        'align': WD_ALIGN_PARAGRAPH.CENTER,
    },
    # 关键词标签: 黑体小四号(12pt)
    'kw_label': {
        'font': '黑体', 'size': Pt(12), 'bold': False,
    },
    # 关键词词条: 宋体小四号(12pt)
    'kw_content': {
        'font': '宋体', 'size': Pt(12), 'bold': False,
    },
    # Keywords标签: Times New Roman加粗小四号(12pt)
    'en_kw_label': {
        'font': 'Times New Roman', 'size': Pt(12), 'bold': True,
    },
    # Keywords词条: Times New Roman小四号(12pt)
    'en_kw_content': {
        'font': 'Times New Roman', 'size': Pt(12), 'bold': False,
    },
    # 致谢标题: 宋体四号加粗(14pt), 居中
    'thanks_title': {
        'font': '宋体', 'size': Pt(14), 'bold': True,
        'align': WD_ALIGN_PARAGRAPH.CENTER,
    },
    # 附录标题: 宋体加粗小二号(18pt), 居中 (与章标题同级)
    'appendix_title': {
        'font': '宋体', 'size': Pt(18), 'bold': True,
        'align': WD_ALIGN_PARAGRAPH.CENTER,
        'before': Pt(6), 'after': Pt(6),
    },
    # 表内文字: 中文宋体五号/英文TNR五号(10.5pt)
    'table_cell': {
        'font': '宋体', 'size': Pt(10.5), 'bold': False,
    },
}


# ─── 工具函数 ───

def get_text(para):
    """获取段落纯文本，去除首尾空白"""
    return para.text.strip()


def get_first_run(para):
    """获取段落中第一个有实质内容的run"""
    for run in para.runs:
        if run.text.strip():
            return run
    return para.runs[0] if para.runs else None


def match_h1(text):
    """匹配章标题: 第X章 或 第X章 xxx"""
    return bool(re.match(r'^第[一二三四五六七八九十\d]+章\s', text))


def match_h2(text):
    """匹配节标题: X.Y 后跟空格 (但不匹配三级及以上的条标题)"""
    return bool(re.match(r'^[\d]+\.[\d]+(\.[\d]+)?\s', text))


def match_h3(text):
    """匹配条标题: X.Y.Z (三级及以上编号)"""
    return bool(re.match(r'^[\d]+\.[\d]+\.[\d]+', text))


def match_h4(text):
    """匹配款/项标题: (1) (一) 等"""
    return bool(re.match(r'^[（(][\d一二三四五六七八九十]+[）)]', text))


def match_ref(text):
    """匹配参考文献条目: [数字] 开头"""
    return bool(re.match(r'^\[\d+\]', text))


def is_ref_title(text):
    """判断是否为参考文献标题"""
    return text in ('参考文献', '参 考 文 献')


def is_abstract_title(text):
    """判断是否为中文摘要标题"""
    return text in ('摘 要', '摘要')


def is_en_abstract_title(text):
    """判断是否为英文摘要标题"""
    return text.strip().upper() == 'ABSTRACT'


def is_thanks_title(text):
    """判断是否为致谢标题"""
    return text in ('致 谢', '致谢')


def match_appendix(text):
    """匹配附录标题: 附录A 或 附录A："""
    return bool(re.match(r'^附录[A-Z]', text))


def match_thanks_title(text):
    """匹配致谢标题"""
    return text in ('致 谢', '致谢')


def set_run_font(run, font_name, font_name_east_asian=None, font_name_ascii=None):
    """
    同时设置西文字体(run.font.name)和东亚字体(eastAsia属性)。
    在中文文档中，必须同时设置两者才能保证中文字体生效。
    """
    if font_name_east_asian is None:
        font_name_east_asian = font_name
    if font_name_ascii is None:
        font_name_ascii = font_name

    run.font.name = font_name

    # 通过XML设置东亚字体
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name_east_asian)
    # 同时设置西文字体
    rFonts.set(qn('w:ascii'), font_name_ascii)
    rFonts.set(qn('w:hAnsi'), font_name_ascii)


def apply_font_spec(para, spec):
    """
    将字体规范应用到段落的所有run上。
    自动处理中文字体名和东亚字体属性。
    """
    font_name = spec['font']
    for run in para.runs:
        if run.text.strip():
            set_run_font(run, font_name, font_name)
        run.font.size = spec['size']
        run.bold = spec.get('bold', False)
    return 1


def apply_paragraph_format(para, spec):
    """
    应用段落格式规范。
    包括对齐、行距、段前段后、首行缩进等。
    """
    pf = para.paragraph_format

    # 对齐方式
    if 'align' in spec and spec['align'] is not None:
        pf.alignment = spec['align']

    # 段前段后
    if 'before' in spec:
        pf.space_before = spec['before']
    if 'after' in spec:
        pf.space_after = spec['after']

    # 行距 (1.25倍)
    if 'line_spacing' in spec:
        pf.line_spacing = spec['line_spacing']

    # 首行缩进
    if 'first_line_indent' in spec:
        pf.first_line_indent = spec['first_line_indent']
    else:
        # 清除首行缩进 (顶格)
        pf.first_line_indent = Pt(0)


def copy_run_format(src_run, dst_run):
    """将一个run的格式复制到另一个run"""
    dst_run.bold = src_run.bold
    dst_run.italic = src_run.italic
    dst_run.underline = src_run.underline
    if src_run.font.size:
        dst_run.font.size = src_run.font.size
    if src_run.font.name:
        dst_run.font.name = src_run.font.name
    if src_run.font.color and src_run.font.color.rgb:
        dst_run.font.color.rgb = src_run.font.color.rgb

    # 复制东亚字体
    src_rPr = src_run._element.find(qn('w:rPr'))
    if src_rPr is not None:
        src_rFonts = src_rPr.find(qn('w:rFonts'))
        if src_rFonts is not None:
            dst_rPr = dst_run._element.get_or_add_rPr()
            dst_rFonts = dst_rPr.find(qn('w:rFonts'))
            if dst_rFonts is None:
                dst_rFonts = OxmlElement('w:rFonts')
                dst_rPr.insert(0, dst_rFonts)
            east = src_rFonts.get(qn('w:eastAsia'))
            if east:
                dst_rFonts.set(qn('w:eastAsia'), east)


# ─── 修复器类 ───

class ThesisFixer:
    """湖南科技大学毕业论文格式修复器"""

    def __init__(self, filepath):
        self.filepath = filepath
        self.doc = Document(filepath)
        self.stats = {
            'page_setup': 0,
            'header': 0,
            'h1': 0,
            'h2': 0,
            'h3': 0,
            'h4': 0,
            'body': 0,
            'ref_title': 0,
            'ref_entry': 0,
            'abstract_title': 0,
            'en_abstract_title': 0,
            'kw_label': 0,
            'kw_content': 0,
            'en_kw_label': 0,
            'en_kw_content': 0,
            'citation': 0,
            'total_paragraphs': 0,
            'errors': [],
        }

        # 检测封面/扉页区域边界（第一个出现"摘要"/"ABSTRACT"/"目录"/"第X章"的段落）
        self._cover_end_idx = 0
        for i, para in enumerate(self.doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            if re.match(r'^摘\s*要$', text) or text.upper() == 'ABSTRACT' \
               or text in ('目 录', '目录') or re.match(r'^目\s+录$', text) \
               or re.match(r'^第[一二三四五六七八九十\d]+章', text):
                self._cover_end_idx = i
                break

    # ── 4.1.1 页面设置修复 ──
    def fix_page_setup(self):
        """修复所有节的页面设置为A4纸，上30mm 下25mm 左25mm 右25mm，页眉20mm 页脚15mm"""
        fixed_count = 0
        for sec in self.doc.sections:
            sec.page_width = PAGE_SPEC['width']
            sec.page_height = PAGE_SPEC['height']
            sec.top_margin = PAGE_SPEC['top']
            sec.bottom_margin = PAGE_SPEC['bottom']
            sec.left_margin = PAGE_SPEC['left']
            sec.right_margin = PAGE_SPEC['right']
            sec.header_distance = PAGE_SPEC['header']
            sec.footer_distance = PAGE_SPEC['footer']
            fixed_count += 1

        self.stats['page_setup'] = fixed_count
        return fixed_count

    # ── 4.1.2 页眉修复 ──
    def fix_header(self):
        """
        从目录页开始设置页眉。
        内容: "湖南科技大学本科生毕业设计（论文）"
        格式: 五号宋体(10.5pt)居中
        """
        fixed_count = 0

        # 查找目录所在段落，确定从哪个节开始设置页眉
        toc_section_index = 0
        for para in self.doc.paragraphs:
            text = get_text(para)
            if text in ('目 录', '目录'):
                # 查找该段落所在的节
                para_element = para._element
                body = para_element.getparent()
                # 向前搜索该段落属于哪个节
                prev = para_element.getprevious()
                while prev is not None:
                    if prev.tag == qn('w:sectPr'):
                        # 找到节属性，说明该段落属于下一个节
                        toc_section_index = len([s for s in body.findall(qn('w:sectPr')) if list(body).index(s) < list(body).index(prev)])
                        break
                    prev = prev.getprevious()
                break

        # 如果找到了目录，从目录所在节开始设置页眉
        # 否则从第一个节开始（跳过封面节）
        start_section = max(0, toc_section_index)

        for si, sec in enumerate(self.doc.sections):
            if si < start_section:
                continue

            header = sec.header
            # 断开与前一节的链接
            header.is_linked_to_previous = False

            # 清空现有页眉内容
            for para in header.paragraphs:
                for run in para.runs:
                    run._element.getparent().remove(run._element)

            # 使用第一个段落设置页眉
            if header.paragraphs:
                para = header.paragraphs[0]
            else:
                para = header.add_paragraph()

            para.alignment = HEADER_SPEC['align']

            # 添加页眉文本
            run = para.add_run(HEADER_SPEC['text'])
            run.font.size = HEADER_SPEC['size']
            set_run_font(run, HEADER_SPEC['font'])

            fixed_count += 1

        self.stats['header'] = fixed_count
        return fixed_count

    # ── 4.2 段落格式修复 ──
    def fix_paragraphs(self):
        """
        遍历所有段落，根据内容匹配类型并修复格式。
        匹配顺序: 从最具体到最宽泛，避免误匹配。
        """
        for i, para in enumerate(self.doc.paragraphs):
            text = get_text(para)
            if not text:
                continue

            # 跳过封面/扉页区域
            if i < self._cover_end_idx:
                continue

            self.stats['total_paragraphs'] += 1

            # 1. 参考文献标题 (最具体匹配)
            if is_ref_title(text):
                apply_font_spec(para, FONT_SPEC['ref_title'])
                apply_paragraph_format(para, {
                    'align': FONT_SPEC['ref_title']['align'],
                    'before': Pt(0), 'after': Pt(0),
                })
                self.stats['ref_title'] += 1
                continue

            # 2. 参考文献条目: [1] 开头
            if match_ref(text):
                self._fix_ref_entry(para)
                apply_paragraph_format(para, {
                    'align': FONT_SPEC['ref_entry']['align'],
                    'before': Pt(0), 'after': Pt(0),
                })
                self.stats['ref_entry'] += 1
                continue

            # 3. 中文摘要标题
            if is_abstract_title(text):
                apply_font_spec(para, FONT_SPEC['abstract_title'])
                apply_paragraph_format(para, {
                    'align': FONT_SPEC['abstract_title']['align'],
                    'before': Pt(0), 'after': Pt(0),
                })
                self.stats['abstract_title'] += 1
                continue

            # 4. 英文摘要标题
            if is_en_abstract_title(text):
                apply_font_spec(para, FONT_SPEC['en_abstract_title'])
                apply_paragraph_format(para, {
                    'align': FONT_SPEC['en_abstract_title']['align'],
                    'before': Pt(0), 'after': Pt(0),
                })
                self.stats['en_abstract_title'] += 1
                continue

            # 5. 关键词行 (含中文关键词)
            if text.startswith('关键词') or text.startswith('关键词：'):
                self._fix_keywords(para, text)
                continue

            # 6. Keywords行 (英文关键词)
            if text.startswith('Keywords') or text.startswith('KeyWords'):
                self._fix_en_keywords(para, text)
                continue

            # 7. 致谢标题
            if is_thanks_title(text):
                apply_font_spec(para, FONT_SPEC['thanks_title'])
                apply_paragraph_format(para, {
                    'align': FONT_SPEC['thanks_title']['align'],
                    'before': Pt(0), 'after': Pt(6),
                })
                self.stats['thanks_title'] = self.stats.get('thanks_title', 0) + 1
                continue

            # 8. 附录标题
            if match_appendix(text):
                apply_font_spec(para, FONT_SPEC['appendix_title'])
                apply_paragraph_format(para, FONT_SPEC['appendix_title'])
                self.stats['appendix_title'] = self.stats.get('appendix_title', 0) + 1
                continue

            # 7. 章标题: 第X章 ...
            if match_h1(text):
                apply_font_spec(para, FONT_SPEC['h1'])
                apply_paragraph_format(para, FONT_SPEC['h1'])
                self.stats['h1'] += 1
                continue

            # 8. 条标题: X.Y.Z ... (三级及以上，先于节标题匹配，避免被节标题误匹配)
            if match_h3(text):
                apply_font_spec(para, FONT_SPEC['h3'])
                apply_paragraph_format(para, FONT_SPEC['h3'])
                self.stats['h3'] += 1
                continue

            # 9. 节标题: X.Y ... (二级，也可能匹配带有空格的X.Y.Z，但条标题已优先处理)
            if match_h2(text):
                apply_font_spec(para, FONT_SPEC['h2'])
                apply_paragraph_format(para, FONT_SPEC['h2'])
                self.stats['h2'] += 1
                continue

            # 10. 款/项标题: (1) (一) 等
            if match_h4(text):
                apply_font_spec(para, FONT_SPEC['h4'])
                apply_paragraph_format(para, FONT_SPEC['h4'])
                self.stats['h4'] += 1
                continue

            # 11. 正文内容 (默认)
            apply_font_spec(para, FONT_SPEC['body'])
            apply_paragraph_format(para, FONT_SPEC['body'])
            self.stats['body'] += 1

    def _fix_ref_entry(self, para):
        """
        修复参考文献条目格式。
        中文字体: 宋体五号(10.5pt)
        西文字体: Times New Roman五号(10.5pt)
        """
        for run in para.runs:
            if not run.text.strip():
                continue
            # 设置字号
            run.font.size = FONT_SPEC['ref_entry']['size']
            run.bold = FONT_SPEC['ref_entry']['bold']
            # 同时设置中文字体为宋体，西文字体为Times New Roman
            set_run_font(run, '宋体', '宋体', 'Times New Roman')

    def _fix_keywords(self, para, text):
        """
        修复中文关键词行格式。
        "关键词"标签: 黑体小四号(12pt)
        词条: 宋体小四号(12pt)
        """
        for run in para.runs:
            if '关键词' in run.text:
                set_run_font(run, FONT_SPEC['kw_label']['font'])
                run.font.size = FONT_SPEC['kw_label']['size']
                run.bold = FONT_SPEC['kw_label']['bold']
                self.stats['kw_label'] += 1
            else:
                set_run_font(run, FONT_SPEC['kw_content']['font'])
                run.font.size = FONT_SPEC['kw_content']['size']
                run.bold = FONT_SPEC['kw_content']['bold']
                self.stats['kw_content'] += 1

    def _fix_en_keywords(self, para, text):
        """
        修复英文关键词行格式。
        "Keywords"标签: Times New Roman加粗小四号(12pt)
        词条: Times New Roman小四号(12pt)
        """
        for run in para.runs:
            run_text = run.text.strip()
            if run_text.upper().startswith('KEYWORD'):
                set_run_font(run, FONT_SPEC['en_kw_label']['font'])
                run.font.size = FONT_SPEC['en_kw_label']['size']
                run.bold = FONT_SPEC['en_kw_label']['bold']
                self.stats['en_kw_label'] += 1
            else:
                set_run_font(run, FONT_SPEC['en_kw_content']['font'])
                run.font.size = FONT_SPEC['en_kw_content']['size']
                run.bold = FONT_SPEC['en_kw_content']['bold']
                self.stats['en_kw_content'] += 1

    # ── 3.3 引用标注修复 ──
    def fix_citations(self):
        """
        将正文中的引用标注 [1], [2,3], [4-6] 等设为上标格式。
        不处理"由文献[8]"这种直接说明的情况。
        """
        total_fixed = 0

        for i, para in enumerate(self.doc.paragraphs):
            text = get_text(para)
            if not text:
                continue
            # 跳过封面/扉页区域
            if i < self._cover_end_idx:
                continue

            # 跳过标题段落
            if match_h1(text) or match_h2(text) or match_h3(text) or match_h4(text):
                continue
            if is_ref_title(text) or is_abstract_title(text) or is_en_abstract_title(text):
                continue
            if match_ref(text):
                continue

            # 检查每个run中的引用标注
            for run in para.runs:
                run_text = run.text
                if not run_text:
                    continue

                # 查找引用标注模式: [数字] 或 [数字,数字] 或 [数字-数字]
                pattern = r'\[\d+(?:[,，\-—]\d+)*\]'
                matches = list(re.finditer(pattern, run_text))

                if not matches:
                    continue

                # 过滤掉"由文献[8]"这类直接说明的情况
                valid_matches = []
                for m in matches:
                    # 检查引用前5个字符是否包含"文献"、"参考"、"见"等
                    before = run_text[max(0, m.start() - 5):m.start()]
                    if not re.search(r'文献|参考|见', before):
                        valid_matches.append(m)

                if not valid_matches:
                    continue

                # 需要拆分run来单独设置上标
                # 收集run的格式信息
                run_bold = run.bold
                run_italic = run.italic
                run_size = run.font.size
                run_font_name = run.font.name
                run_color = run.font.color.rgb if run.font.color and run.font.color.rgb else None

                # 获取东亚字体
                east_asian = None
                rPr = run._element.find(qn('w:rPr'))
                if rPr is not None:
                    rFonts = rPr.find(qn('w:rFonts'))
                    if rFonts is not None:
                        east_asian = rFonts.get(qn('w:eastAsia'))

                # 保存run的XML元素引用，用于后续插入新run
                run_element = run._element
                parent = run_element.getparent()

                # 在原run的文本中，将引用标注替换为占位标记
                # 并记录需要插入的引用标注位置
                pieces = []
                last_end = 0
                for m in valid_matches:
                    if m.start() > last_end:
                        pieces.append(('normal', run_text[last_end:m.start()]))
                    pieces.append(('superscript', m.group()))
                    last_end = m.end()

                if last_end < len(run_text):
                    pieces.append(('normal', run_text[last_end:]))

                # 修改原run的文本为第一个piece
                run.text = pieces[0][1]

                # 为后续pieces创建新run，插入到原run之后
                insert_after = run_element
                for piece_type, piece_text in pieces[1:]:
                    new_run_element = OxmlElement('w:r')
                    # 复制rPr
                    new_rPr = OxmlElement('w:rPr')
                    if run_bold:
                        new_rPr_b = OxmlElement('w:b')
                        new_rPr.append(new_rPr_b)
                    if run_italic:
                        new_rPr_i = OxmlElement('w:i')
                        new_rPr.append(new_rPr_i)
                    if run_size:
                        new_rPr_sz = OxmlElement('w:sz')
                        new_rPr_sz.set(qn('w:val'), str(int(run_size)))
                        new_rPr.append(new_rPr_sz)
                    if run_font_name or east_asian:
                        new_rPr_rFonts = OxmlElement('w:rFonts')
                        if run_font_name:
                            new_rPr_rFonts.set(qn('w:ascii'), run_font_name)
                            new_rPr_rFonts.set(qn('w:hAnsi'), run_font_name)
                        if east_asian:
                            new_rPr_rFonts.set(qn('w:eastAsia'), east_asian)
                        new_rPr.append(new_rPr_rFonts)
                    if run_color:
                        new_rPr_color = OxmlElement('w:color')
                        new_rPr_color.set(qn('w:val'), str(run_color))
                        new_rPr.append(new_rPr_color)

                    # 上标设置
                    if piece_type == 'superscript':
                        new_rPr_vertAlign = OxmlElement('w:vertAlign')
                        new_rPr_vertAlign.set(qn('w:val'), 'superscript')
                        new_rPr.append(new_rPr_vertAlign)
                        total_fixed += 1

                    new_run_element.append(new_rPr)

                    # 添加文本
                    new_t = OxmlElement('w:t')
                    new_t.set(qn('xml:space'), 'preserve')
                    new_t.text = piece_text
                    new_run_element.append(new_t)

                    # 插入到原run之后
                    insert_after.addnext(new_run_element)
                    insert_after = new_run_element

        self.stats['citation'] = total_fixed
        return total_fixed

    # ── 表内文字格式修复 ──
    def fix_table_cells(self):
        """修复表内文字格式：中文五号宋体，英文五号Times New Roman"""
        fixed_count = 0
        for ti, table in enumerate(self.doc.tables):
            # 跳过封面表格（第一个表格，通常包含"题目"/"作者"等）
            if ti == 0:
                first_cell = table.rows[0].cells[0].text.strip() if table.rows else ''
                if first_cell in ('题目', '作者', '学院', '专业', '学号', '指导教师'):
                    continue
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if not run.text.strip():
                                continue
                            run_text = run.text.strip()
                            has_chinese = bool(re.search(r'[\u4e00-\u9fff]', run_text))
                            run.font.size = FONT_SPEC['table_cell']['size']
                            if has_chinese:
                                set_run_font(run, FONT_SPEC['table_cell']['font'])
                            else:
                                set_run_font(run, 'Times New Roman')
                            fixed_count += 1
        self.stats['table_cell'] = fixed_count
        return fixed_count

    # ── 分页符插入 ──
    def fix_page_breaks(self):
        """在每章、参考文献、致谢、附录之前插入分页符"""
        fixed_count = 0
        paragraphs = list(self.doc.paragraphs)
        
        for i, para in enumerate(paragraphs):
            text = get_text(para)
            if not text:
                continue
            # 跳过封面/扉页区域
            if i < self._cover_end_idx:
                continue
            
            # 检查是否需要插入分页符的标题
            needs_break = (
                match_h1(text) or  # 章标题
                text in ('参考文献', '参 考 文 献') or
                match_thanks_title(text) or
                match_appendix(text)
            )
            
            if needs_break and i > 0:
                # 检查前一个段落是否已经有分页符
                prev_para = paragraphs[i - 1]
                has_page_break = False
                
                # 检查前一段落的pPr中是否有pageBreakBefore
                pPr = prev_para._element.find(qn('w:pPr'))
                if pPr is not None:
                    page_break = pPr.find(qn('w:pageBreakBefore'))
                    if page_break is not None:
                        has_page_break = True
                
                # 检查前一段落中是否有分页符
                if not has_page_break:
                    for run_elem in prev_para._element.findall(qn('w:r')):
                        for br in run_elem.findall(qn('w:br')):
                            if br.get(qn('w:type')) == 'page':
                                has_page_break = True
                                break
                        if has_page_break:
                            break
                
                # 如果没有分页符，则在当前段落前插入
                if not has_page_break:
                    # 在当前段落的pPr中添加pageBreakBefore
                    pPr = para._element.find(qn('w:pPr'))
                    if pPr is None:
                        pPr = OxmlElement('w:pPr')
                        para._element.insert(0, pPr)
                    
                    page_break = OxmlElement('w:pageBreakBefore')
                    pPr.append(page_break)
                    fixed_count += 1
        
        self.stats['page_break'] = fixed_count
        return fixed_count

    # ── 汇总 ──
    def run_all(self):
        """执行所有修复项并输出统计"""
        print(f"\n{'=' * 60}")
        print(f"  湖南科技大学毕业论文格式修复")
        print(f"  文件: {self.filepath}")
        print(f"{'=' * 60}")

        # 1. 页面设置
        self.fix_page_setup()
        print(f"\n  [4.1.1] 页面设置: 修复 {self.stats['page_setup']} 个节")

        # 2. 页眉
        self.fix_header()
        print(f"  [4.1.2] 页眉: 修复 {self.stats['header']} 个节")

        # 3. 段落格式
        self.fix_paragraphs()
        print(f"\n  [4.2] 段落格式修复:")
        print(f"    章标题 (第X章):                     {self.stats['h1']} 个")
        print(f"    节标题 (X.Y):                        {self.stats['h2']} 个")
        print(f"    条标题 (X.Y.Z):                      {self.stats['h3']} 个")
        print(f"    款/项标题 ((一)/(1)):                 {self.stats['h4']} 个")
        print(f"    正文内容:                            {self.stats['body']} 个")
        print(f"    参考文献标题:                        {self.stats['ref_title']} 个")
        print(f"    参考文献条目:                        {self.stats['ref_entry']} 个")
        print(f"    摘要标题:                            {self.stats['abstract_title']} 个")
        print(f"    ABSTRACT标题:                        {self.stats['en_abstract_title']} 个")
        print(f"    关键词标签:                          {self.stats['kw_label']} 个")
        print(f"    关键词词条:                          {self.stats['kw_content']} 个")
        print(f"    Keywords标签:                        {self.stats['en_kw_label']} 个")
        print(f"    Keywords词条:                        {self.stats['en_kw_content']} 个")

        # 4. 引用标注
        self.fix_citations()
        print(f"\n  [3.3] 引用标注上标: 修复 {self.stats['citation']} 处")

        # 5. 表内文字
        self.fix_table_cells()
        print(f"  [4.5] 表内文字格式: 修复 {self.stats.get('table_cell', 0)} 处")

        # 6. 分页符
        self.fix_page_breaks()
        print(f"  [4.3] 分页符: 插入 {self.stats.get('page_break', 0)} 处")

        # 总计
        total_fixed = (
            self.stats['page_setup'] + self.stats['header'] +
            self.stats['h1'] + self.stats['h2'] + self.stats['h3'] + self.stats['h4'] +
            self.stats['body'] + self.stats['ref_title'] + self.stats['ref_entry'] +
            self.stats['abstract_title'] + self.stats['en_abstract_title'] +
            self.stats['kw_label'] + self.stats['kw_content'] +
            self.stats['en_kw_label'] + self.stats['en_kw_content'] +
            self.stats['citation'] + self.stats.get('table_cell', 0) +
            self.stats.get('thanks_title', 0) + self.stats.get('appendix_title', 0) +
            self.stats.get('page_break', 0)
        )

        print(f"\n{'─' * 60}")
        print(f"  总计修复: {total_fixed} 项")
        print(f"  处理段落总数: {self.stats['total_paragraphs']} 个")
        if self.stats['errors']:
            print(f"  修复中出现的问题: {len(self.stats['errors'])} 个")
            for err in self.stats['errors']:
                print(f"    ! {err}")
        print(f"{'─' * 60}")

        return total_fixed

    def save(self, output_path=None):
        """保存修复后的文档"""
        if output_path is None:
            # 默认在原文件名后加 _fixed
            base, ext = os.path.splitext(self.filepath)
            output_path = f"{base}_fixed{ext}"

        self.doc.save(output_path)
        print(f"\n  已保存修复后的文档: {output_path}")
        return output_path


# ─── 主入口 ───

def main():
    if len(sys.argv) < 2:
        print("用法: python fix_thesis.py <docx文件路径> [输出路径]")
        print("示例: python fix_thesis.py 论文.docx")
        print("      python fix_thesis.py 论文.docx 论文_修复.docx")
        sys.exit(1)

    filepath = sys.argv[1]
    if not os.path.exists(filepath):
        print(f"错误: 文件不存在 - {filepath}")
        sys.exit(1)

    if not filepath.lower().endswith('.docx'):
        print(f"错误: 仅支持 .docx 格式文件 - {filepath}")
        sys.exit(1)

    output_path = sys.argv[2] if len(sys.argv) >= 3 else None

    fixer = ThesisFixer(filepath)
    fixer.run_all()
    fixer.save(output_path)


if __name__ == '__main__':
    main()