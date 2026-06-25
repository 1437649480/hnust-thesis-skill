# -*- coding: utf-8 -*-
"""
湖南科技大学毕业论文格式验证脚本
逐项验证修复后的文档格式是否完全符合规范要求。
用法: python verify_thesis.py <docx文件路径>
"""

import sys
import os
import re
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ─── 常量定义 ───
# 容差: 50000 EMU (约0.139mm)
TOLERANCE = 50000

# 规范要求的页面设置
PAGE_SPEC = {
    'width': Cm(21),        # A4宽度 210mm
    'height': Cm(29.7),     # A4高度 297mm
    'top': Cm(3),           # 上页边距 30mm
    'bottom': Cm(2.5),      # 下页边距 25mm
    'left': Cm(2.5),        # 左页边距 25mm
    'right': Cm(2.5),       # 右页边距 25mm
    'header': Cm(2),        # 页眉距边界 20mm
    'footer': Cm(1.5),      # 页脚距边界 15mm
}

# 页眉规范
HEADER_SPEC = {
    'text': '湖南科技大学本科生毕业设计（论文）',
    'font': '宋体',
    'size': Pt(10.5),       # 五号
    'align': WD_ALIGN_PARAGRAPH.CENTER,
}

# 页码规范
PAGENUM_SPEC = {
    'font': 'Times New Roman',
    'size': Pt(10.5),       # 五号
    'align': WD_ALIGN_PARAGRAPH.CENTER,
}

# 字体字号规范
FONT_SPEC = {
    'h1': {  # 章标题: 第X章
        'font': '宋体', 'size': Pt(18), 'bold': True,
        'align': WD_ALIGN_PARAGRAPH.CENTER,
        'before': Pt(6), 'after': Pt(6),
    },
    'h2': {  # 节标题: X.Y
        'font': '宋体', 'size': Pt(14), 'bold': True,
        'before': Pt(6), 'after': Pt(6),
    },
    'h3': {  # 条标题: X.Y.Z
        'font': '宋体', 'size': Pt(12), 'bold': True,
        'before': Pt(6), 'after': Pt(6),
    },
    'body': {  # 正文
        'font': '宋体', 'size': Pt(12), 'bold': False,
        'line_spacing': 1.25, 'first_line_indent': Pt(24),
    },
    'ref_title': {  # 参考文献标题
        'font': '宋体', 'size': Pt(14), 'bold': True,
        'align': WD_ALIGN_PARAGRAPH.CENTER,
    },
    'ref_entry': {  # 参考文献条目
        'font': '宋体', 'size': Pt(10.5), 'bold': False,
    },
}


# ─── 辅助函数 ───
def get_text(para):
    """获取段落纯文本，去除首尾空白"""
    return para.text.strip()


def get_first_run(para):
    """获取段落中第一个有实质内容的run"""
    for run in para.runs:
        if run.text.strip():
            return run
    return para.runs[0] if para.runs else None


def emu_to_mm(emu):
    """将EMU转换为毫米"""
    return emu / 360000 * 10  # 1 cm = 360000 EMU, 1 cm = 10 mm


def emu_to_cm(emu):
    """将EMU转换为厘米"""
    return emu / 360000


def pt_value(pt):
    """获取Pt对象的数值（整数EMU）"""
    if pt is None:
        return None
    return int(pt)


def check_emu(actual, expected, tolerance=TOLERANCE):
    """检查EMU值是否在容差范围内"""
    if actual is None:
        return False
    return abs(int(actual) - int(expected)) <= tolerance


def match_h1(text):
    """匹配章标题: 第X章 或 第X章 xxx"""
    return bool(re.match(r'第[一二三四五六七八九十\d]+章', text))


def match_h2(text):
    """匹配节标题: X.Y 后跟空格"""
    return bool(re.match(r'\d+\.\d+\s', text))


def match_h3(text):
    """匹配条标题: X.Y.Z 开头（排除节标题）"""
    return bool(re.match(r'\d+\.\d+\.\d+', text))


def match_ref(text):
    """匹配参考文献条目: [数字] 开头"""
    return bool(re.match(r'^\[\d+\]', text))


def has_page_field(paragraph):
    """检查段落中是否包含PAGE域"""
    # PAGE域在XML中的结构: fldChar(begin) -> instrText(PAGE) -> fldChar(separate) -> 结果 -> fldChar(end)
    runs = paragraph._element.findall('.//' + qn('w:r'))
    has_begin = False
    has_instr_page = False
    for r in runs:
        fld_chars = r.findall(qn('w:fldChar'))
        for fc in fld_chars:
            if fc.get(qn('w:fldCharType')) == 'begin':
                has_begin = True
        instr_texts = r.findall(qn('w:instrText'))
        for it in instr_texts:
            if it.text and 'PAGE' in it.text.upper():
                has_instr_page = True
    return has_begin and has_instr_page


def has_page_field_in_footer(footer):
    """检查页脚中是否包含PAGE域"""
    for para in footer.paragraphs:
        if has_page_field(para):
            return True
    return False


def check_footer_dash_format(footer):
    """检查页脚文本是否包含 '-' 前后缀格式 (如 -1- 或 -i-)"""
    for para in footer.paragraphs:
        text = para.text.strip()
        # 匹配 -数字- 或 -罗马数字- 格式
        if re.match(r'^-.+-$', text):
            return True
    return False


# ─── 验证器类 ───
class ThesisVerifier:
    """湖南科技大学毕业论文格式验证器"""

    def __init__(self, filepath):
        self.filepath = filepath
        self.doc = Document(filepath)
        self.sections = self.doc.sections
        self.paragraphs = self.doc.paragraphs

        # 统计
        self.total = 0
        self.passed = 0
        self.failed = 0

        # 存储所有结果: [(category, ok, label, detail), ...]
        self.results = []

    def verify(self, category, condition, label, detail=None):
        """记录一项验证结果"""
        self.total += 1
        if condition:
            self.passed += 1
            self.results.append((category, True, label, None))
        else:
            self.failed += 1
            self.results.append((category, False, label, detail))

    # ══════════════════════════════════════════════════════════
    # 页面设置验证
    # ══════════════════════════════════════════════════════════
    def verify_page_setup(self):
        """验证页面设置: A4纸张、页边距、页眉页脚距离"""
        sec = self.sections[0]

        # A4纸张 210mm×297mm
        w_ok = check_emu(sec.page_width, PAGE_SPEC['width'])
        h_ok = check_emu(sec.page_height, PAGE_SPEC['height'])
        if w_ok and h_ok:
            self.verify('页面', True, 'A4纸张 210mm×297mm')
        else:
            actual_w = emu_to_mm(sec.page_width)
            actual_h = emu_to_mm(sec.page_height)
            self.verify('页面', False, 'A4纸张 210mm×297mm',
                        f'实际: {actual_w:.1f}mm×{actual_h:.1f}mm')

        # 上页边距 30mm
        ok = check_emu(sec.top_margin, PAGE_SPEC['top'])
        actual = emu_to_mm(sec.top_margin)
        self.verify('页面', ok, '上页边距应为30mm',
                    f'实际: {actual:.1f}mm' if not ok else None)

        # 下页边距 25mm
        ok = check_emu(sec.bottom_margin, PAGE_SPEC['bottom'])
        actual = emu_to_mm(sec.bottom_margin)
        self.verify('页面', ok, '下页边距应为25mm',
                    f'实际: {actual:.1f}mm' if not ok else None)

        # 左页边距 25mm
        ok = check_emu(sec.left_margin, PAGE_SPEC['left'])
        actual = emu_to_mm(sec.left_margin)
        self.verify('页面', ok, '左页边距应为25mm',
                    f'实际: {actual:.1f}mm' if not ok else None)

        # 右页边距 25mm
        ok = check_emu(sec.right_margin, PAGE_SPEC['right'])
        actual = emu_to_mm(sec.right_margin)
        self.verify('页面', ok, '右页边距应为25mm',
                    f'实际: {actual:.1f}mm' if not ok else None)

        # 页眉距离 20mm
        ok = check_emu(sec.header_distance, PAGE_SPEC['header'])
        actual = emu_to_mm(sec.header_distance)
        self.verify('页面', ok, '页眉距边界应为20mm',
                    f'实际: {actual:.1f}mm' if not ok else None)

        # 页脚距离 15mm
        ok = check_emu(sec.footer_distance, PAGE_SPEC['footer'])
        actual = emu_to_mm(sec.footer_distance)
        self.verify('页面', ok, '页脚距边界应为15mm',
                    f'实际: {actual:.1f}mm' if not ok else None)

    # ══════════════════════════════════════════════════════════
    # 页眉验证
    # ══════════════════════════════════════════════════════════
    def verify_header(self):
        """验证每节页眉: 内容、字体、居中"""
        for si, sec in enumerate(self.sections):
            header = sec.header
            if header.is_linked_to_previous:
                self.verify('页眉', False,
                            f'节{si + 1} 页眉未断开"链接到前一节"')
                continue

            paras = header.paragraphs
            if not paras or not get_text(paras[0]):
                self.verify('页眉', False,
                            f'节{si + 1} 页眉为空')
                continue

            p = paras[0]
            text = get_text(p)

            # 页眉内容
            if HEADER_SPEC['text'] in text:
                self.verify('页眉', True,
                            f'节{si + 1} 页眉内容含"{HEADER_SPEC["text"]}"')
            else:
                self.verify('页眉', False,
                            f'节{si + 1} 页眉内容应为"{HEADER_SPEC["text"]}"',
                            f'实际: "{text[:40]}"')

            # 页眉居中
            if p.alignment == HEADER_SPEC['align']:
                self.verify('页眉', True, f'节{si + 1} 页眉居中')
            else:
                self.verify('页眉', False, f'节{si + 1} 页眉应居中',
                            f'实际: {p.alignment}')

            # 页眉字体
            run = get_first_run(p)
            if run:
                # 字体
                if run.font.name == HEADER_SPEC['font']:
                    self.verify('页眉', True,
                                f'节{si + 1} 页眉字体为五号宋体(10.5pt)')
                else:
                    self.verify('页眉', False,
                                f'节{si + 1} 页眉字体应为宋体',
                                f'实际: {run.font.name}')

                # 字号
                if run.font.size and check_emu(run.font.size, HEADER_SPEC['size']):
                    self.verify('页眉', True,
                                f'节{si + 1} 页眉字号为五号(10.5pt)')
                else:
                    actual_pt = run.font.size / 12700 if run.font.size else None
                    self.verify('页眉', False,
                                f'节{si + 1} 页眉字号应为五号(10.5pt)',
                                f'实际: {actual_pt}pt')
            else:
                self.verify('页眉', False,
                            f'节{si + 1} 页眉无可读取的字体信息')

    # ══════════════════════════════════════════════════════════
    # 页码验证
    # ══════════════════════════════════════════════════════════
    def verify_page_number(self):
        """验证页码: 格式、字体、PAGE域、前后缀"""
        num_sections = len(self.sections)

        for si, sec in enumerate(self.sections):
            footer = sec.footer

            # 获取页码格式
            sect_pr = sec._sectPr
            pgnt = sect_pr.find(qn('w:pgNumType'))
            fmt = pgnt.get(qn('w:fmt')) if pgnt is not None else None
            start = pgnt.get(qn('w:start')) if pgnt is not None else None

            if num_sections > 1:
                # 多节: 前置部分(节1)罗马数字, 正文(节2)阿拉伯数字
                if si == 0:
                    # 前置部分用罗马数字
                    if fmt == 'lowerRoman':
                        self.verify('页码', True,
                                    f'节{si + 1} 前置部分页码为罗马数字(lowerRoman)')
                    else:
                        self.verify('页码', False,
                                    f'节{si + 1} 前置部分页码应为罗马数字(lowerRoman)',
                                    f'实际: {fmt}')
                else:
                    # 正文部分用阿拉伯数字
                    if fmt == 'decimal':
                        self.verify('页码', True,
                                    f'节{si + 1} 正文页码为阿拉伯数字(decimal)')
                    else:
                        self.verify('页码', False,
                                    f'节{si + 1} 正文页码应为阿拉伯数字(decimal)',
                                    f'实际: {fmt}')
                    # 正文起始页码应为1
                    if start == '1':
                        self.verify('页码', True,
                                    f'节{si + 1} 正文起始页码为1')
                    else:
                        self.verify('页码', False,
                                    f'节{si + 1} 正文起始页码应为1',
                                    f'实际: start={start}')
            else:
                # 单节: 阿拉伯数字
                if fmt == 'decimal':
                    self.verify('页码', True,
                                '页码格式为阿拉伯数字(decimal)')
                else:
                    self.verify('页码', False,
                                '页码格式应为阿拉伯数字(decimal)',
                                f'实际: {fmt}')

            # 检查页脚是否链接到前一节
            if footer.is_linked_to_previous and si > 0:
                self.verify('页码', False,
                            f'节{si + 1} 页脚未断开"链接到前一节"')
                continue

            # 检查页脚段落
            footer_paras = footer.paragraphs
            if not footer_paras:
                self.verify('页码', False,
                            f'节{si + 1} 页脚为空')
                continue

            fp = footer_paras[0]

            # 页脚居中
            if fp.alignment == PAGENUM_SPEC['align']:
                self.verify('页码', True, f'节{si + 1} 页脚居中')
            else:
                self.verify('页码', False, f'节{si + 1} 页脚应居中',
                            f'实际: {fp.alignment}')

            # 页脚字体: 五号Times New Roman
            run = get_first_run(fp)
            if run:
                font_ok = (run.font.name and
                           run.font.name.upper().replace(' ', '') == 'TIMESNEWROMAN')
                size_ok = (run.font.size and
                           check_emu(run.font.size, PAGENUM_SPEC['size']))

                if font_ok and size_ok:
                    self.verify('页码', True,
                                f'节{si + 1} 页脚字体为五号Times New Roman')
                else:
                    issues = []
                    if not font_ok:
                        issues.append(f'字体: {run.font.name}')
                    if not size_ok:
                        issues.append(f'字号: {run.font.size / 12700 if run.font.size else "None"}pt')
                    self.verify('页码', False,
                                f'节{si + 1} 页脚字体应为五号Times New Roman',
                                f'实际: {", ".join(issues)}')
            else:
                self.verify('页码', False,
                            f'节{si + 1} 页脚无可读取的字体信息')

            # 页脚包含PAGE域
            if has_page_field_in_footer(footer):
                self.verify('页码', True,
                            f'节{si + 1} 页脚包含PAGE域')
            else:
                self.verify('页码', False,
                            f'节{si + 1} 页脚应包含PAGE域',
                            '实际: 未检测到PAGE域')

            # 页脚包含 "-" 前后缀
            if check_footer_dash_format(footer):
                self.verify('页码', True,
                            f'节{si + 1} 页脚含"-"前后缀')
            else:
                footer_text = get_text(fp)
                self.verify('页码', False,
                            f'节{si + 1} 页脚应含"-"前后缀',
                            f'实际: "{footer_text[:20]}"')

    # ══════════════════════════════════════════════════════════
    # 字体字号验证
    # ══════════════════════════════════════════════════════════
    def verify_font_and_paragraph(self):
        """验证各级标题和正文的字体字号段落格式"""
        has_h1 = False
        has_h2 = False
        has_h3 = False
        has_body = False
        has_ref_title = False
        has_ref_entry = False

        for para in self.paragraphs:
            text = get_text(para)
            if not text:
                continue

            # ── 章标题: 第X章 ──
            if match_h1(text):
                has_h1 = True
                self._verify_h1(para, text)
                continue

            # ── 节标题: X.Y (后跟空格) ──
            if match_h2(text):
                has_h2 = True
                self._verify_h2(para, text)
                continue

            # ── 条标题: X.Y.Z ──
            if match_h3(text):
                has_h3 = True
                self._verify_h3(para, text)
                continue

            # ── 参考文献标题 ──
            if text in ('参考文献', '参 考 文 献'):
                has_ref_title = True
                self._verify_ref_title(para, text)
                continue

            # ── 参考文献条目: [数字] ──
            if match_ref(text):
                has_ref_entry = True
                self._verify_ref_entry(para, text)
                continue

            # ── 正文 ──
            # 跳过封面、目录、摘要等前置部分的标题
            if text in ('目 录', '目录', '摘 要', '摘要', 'ABSTRACT',
                        '致 谢', '致谢', '附录', '附 录'):
                continue
            # 跳过太短的文本（可能是页眉页脚标记等）
            if len(text) < 4:
                continue

            has_body = True
            self._verify_body(para, text)

        # 汇总: 如果没有检测到某类标题，报告
        if not has_h1:
            self.verify('字体', False, '未检测到章标题(第X章)',
                        '文档中可能缺少章标题')
        if not has_h2:
            self.verify('字体', False, '未检测到节标题(X.Y)',
                        '文档中可能缺少节标题')
        if not has_h3:
            self.verify('字体', False, '未检测到条标题(X.Y.Z)',
                        '文档中可能缺少条标题')
        if not has_ref_title:
            self.verify('字体', False, '未检测到参考文献标题',
                        '文档中可能缺少参考文献标题')
        if not has_ref_entry:
            self.verify('字体', False, '未检测到参考文献条目',
                        '文档中可能缺少参考文献条目')

    def _verify_h1(self, para, text):
        """验证章标题格式"""
        spec = FONT_SPEC['h1']
        run = get_first_run(para)
        short = text[:25]

        if run:
            # 字体
            if run.font.name == spec['font']:
                self.verify('字体', True, f'章标题"{short}" 宋体')
            else:
                self.verify('字体', False, f'章标题"{short}" 字体应为宋体',
                            f'实际: {run.font.name}')

            # 加粗
            if run.bold:
                self.verify('字体', True, f'章标题"{short}" 加粗')
            else:
                self.verify('字体', False, f'章标题"{short}" 应加粗')

            # 字号 小二号 18pt
            if run.font.size and check_emu(run.font.size, spec['size']):
                self.verify('字体', True, f'章标题"{short}" 小二号(18pt)')
            else:
                actual_pt = run.font.size / 12700 if run.font.size else None
                self.verify('字体', False, f'章标题"{short}" 字号应为小二号(18pt)',
                            f'实际: {actual_pt}pt' if actual_pt else '实际: 无字号')

        # 居中
        if para.alignment == spec['align']:
            self.verify('字体', True, f'章标题"{short}" 居中')
        else:
            self.verify('字体', False, f'章标题"{short}" 应居中',
                        f'实际: {para.alignment}')

        # 段前段后各6pt
        pf = para.paragraph_format
        if pf.space_before is not None and check_emu(pf.space_before, spec['before']):
            self.verify('字体', True, f'章标题"{short}" 段前6pt')
        else:
            actual = pf.space_before / 12700 if pf.space_before else 0
            self.verify('字体', False, f'章标题"{short}" 段前应为6pt',
                        f'实际: {actual:.0f}pt')

        if pf.space_after is not None and check_emu(pf.space_after, spec['after']):
            self.verify('字体', True, f'章标题"{short}" 段后6pt')
        else:
            actual = pf.space_after / 12700 if pf.space_after else 0
            self.verify('字体', False, f'章标题"{short}" 段后应为6pt',
                        f'实际: {actual:.0f}pt')

    def _verify_h2(self, para, text):
        """验证节标题格式"""
        spec = FONT_SPEC['h2']
        run = get_first_run(para)
        short = text[:25]

        if run:
            if run.font.name == spec['font']:
                self.verify('字体', True, f'节标题"{short}" 宋体')
            else:
                self.verify('字体', False, f'节标题"{short}" 字体应为宋体',
                            f'实际: {run.font.name}')

            if run.bold:
                self.verify('字体', True, f'节标题"{short}" 加粗')
            else:
                self.verify('字体', False, f'节标题"{short}" 应加粗')

            if run.font.size and check_emu(run.font.size, spec['size']):
                self.verify('字体', True, f'节标题"{short}" 四号(14pt)')
            else:
                actual_pt = run.font.size / 12700 if run.font.size else None
                self.verify('字体', False, f'节标题"{short}" 字号应为四号(14pt)',
                            f'实际: {actual_pt}pt' if actual_pt else '实际: 无字号')

        # 段前段后各6pt
        pf = para.paragraph_format
        if pf.space_before is not None and check_emu(pf.space_before, spec['before']):
            self.verify('字体', True, f'节标题"{short}" 段前6pt')
        else:
            actual = pf.space_before / 12700 if pf.space_before else 0
            self.verify('字体', False, f'节标题"{short}" 段前应为6pt',
                        f'实际: {actual:.0f}pt')

        if pf.space_after is not None and check_emu(pf.space_after, spec['after']):
            self.verify('字体', True, f'节标题"{short}" 段后6pt')
        else:
            actual = pf.space_after / 12700 if pf.space_after else 0
            self.verify('字体', False, f'节标题"{short}" 段后应为6pt',
                        f'实际: {actual:.0f}pt')

    def _verify_h3(self, para, text):
        """验证条标题格式"""
        spec = FONT_SPEC['h3']
        run = get_first_run(para)
        short = text[:25]

        if run:
            if run.font.name == spec['font']:
                self.verify('字体', True, f'条标题"{short}" 宋体')
            else:
                self.verify('字体', False, f'条标题"{short}" 字体应为宋体',
                            f'实际: {run.font.name}')

            if run.bold:
                self.verify('字体', True, f'条标题"{short}" 加粗')
            else:
                self.verify('字体', False, f'条标题"{short}" 应加粗')

            if run.font.size and check_emu(run.font.size, spec['size']):
                self.verify('字体', True, f'条标题"{short}" 小四号(12pt)')
            else:
                actual_pt = run.font.size / 12700 if run.font.size else None
                self.verify('字体', False, f'条标题"{short}" 字号应为小四号(12pt)',
                            f'实际: {actual_pt}pt' if actual_pt else '实际: 无字号')

        # 段前段后各6pt
        pf = para.paragraph_format
        if pf.space_before is not None and check_emu(pf.space_before, spec['before']):
            self.verify('字体', True, f'条标题"{short}" 段前6pt')
        else:
            actual = pf.space_before / 12700 if pf.space_before else 0
            self.verify('字体', False, f'条标题"{short}" 段前应为6pt',
                        f'实际: {actual:.0f}pt')

        if pf.space_after is not None and check_emu(pf.space_after, spec['after']):
            self.verify('字体', True, f'条标题"{short}" 段后6pt')
        else:
            actual = pf.space_after / 12700 if pf.space_after else 0
            self.verify('字体', False, f'条标题"{short}" 段后应为6pt',
                        f'实际: {actual:.0f}pt')

    def _verify_body(self, para, text):
        """验证正文格式: 宋体小四号(12pt), 1.25倍行距, 首行缩进24pt"""
        spec = FONT_SPEC['body']
        run = get_first_run(para)
        short = text[:20]
        pf = para.paragraph_format

        if run:
            # 字体
            if run.font.name == spec['font']:
                self.verify('字体', True, f'正文"{short}" 宋体')
            else:
                self.verify('字体', False, f'正文"{short}" 字体应为宋体',
                            f'实际: {run.font.name}')

            # 字号 小四号 12pt
            if run.font.size and check_emu(run.font.size, spec['size']):
                self.verify('字体', True, f'正文"{short}" 小四号(12pt)')
            else:
                actual_pt = run.font.size / 12700 if run.font.size else None
                self.verify('字体', False, f'正文"{short}" 字号应为小四号(12pt)',
                            f'实际: {actual_pt}pt' if actual_pt else '实际: 无字号')

        # 1.25倍行距
        if pf.line_spacing is not None:
            if abs(pf.line_spacing - spec['line_spacing']) <= 0.05:
                self.verify('字体', True, f'正文"{short}" 1.25倍行距')
            else:
                self.verify('字体', False, f'正文"{short}" 行距应为1.25倍',
                            f'实际: {pf.line_spacing}倍')
        else:
            self.verify('字体', False, f'正文"{short}" 行距应为1.25倍',
                        '实际: 未设置行距')

        # 首行缩进 24pt
        if pf.first_line_indent is not None:
            if check_emu(pf.first_line_indent, spec['first_line_indent']):
                self.verify('字体', True, f'正文"{short}" 首行缩进24pt')
            else:
                actual_pt = pf.first_line_indent / 12700
                self.verify('字体', False, f'正文"{short}" 首行缩进应为24pt',
                            f'实际: {actual_pt:.0f}pt')
        else:
            self.verify('字体', False, f'正文"{short}" 应有首行缩进',
                        '实际: 无首行缩进')

    def _verify_ref_title(self, para, text):
        """验证参考文献标题: 宋体四号加粗(14pt), 居中"""
        spec = FONT_SPEC['ref_title']
        run = get_first_run(para)

        if run:
            if run.font.name == spec['font']:
                self.verify('字体', True, '参考文献标题 宋体')
            else:
                self.verify('字体', False, '参考文献标题 字体应为宋体',
                            f'实际: {run.font.name}')

            if run.bold:
                self.verify('字体', True, '参考文献标题 加粗')
            else:
                self.verify('字体', False, '参考文献标题 应加粗')

            if run.font.size and check_emu(run.font.size, spec['size']):
                self.verify('字体', True, '参考文献标题 四号(14pt)')
            else:
                actual_pt = run.font.size / 12700 if run.font.size else None
                self.verify('字体', False, '参考文献标题 字号应为四号(14pt)',
                            f'实际: {actual_pt}pt' if actual_pt else '实际: 无字号')

        if para.alignment == spec['align']:
            self.verify('字体', True, '参考文献标题 居中')
        else:
            self.verify('字体', False, '参考文献标题 应居中',
                        f'实际: {para.alignment}')

    def _verify_ref_entry(self, para, text):
        """验证参考文献条目: 顶格, 五号宋体(10.5pt)"""
        spec = FONT_SPEC['ref_entry']
        run = get_first_run(para)
        short = text[:30]

        if run:
            if run.font.name == spec['font']:
                self.verify('字体', True, f'参考文献条目"{short}" 宋体')
            else:
                self.verify('字体', False, f'参考文献条目"{short}" 字体应为宋体',
                            f'实际: {run.font.name}')

            if run.font.size and check_emu(run.font.size, spec['size']):
                self.verify('字体', True, f'参考文献条目"{short}" 五号(10.5pt)')
            else:
                actual_pt = run.font.size / 12700 if run.font.size else None
                self.verify('字体', False, f'参考文献条目"{short}" 字号应为五号(10.5pt)',
                            f'实际: {actual_pt}pt' if actual_pt else '实际: 无字号')

        # 顶格: 无首行缩进
        pf = para.paragraph_format
        if pf.first_line_indent is None or pf.first_line_indent == 0:
            self.verify('字体', True, f'参考文献条目"{short}" 顶格')
        else:
            actual_pt = pf.first_line_indent / 12700
            self.verify('字体', False, f'参考文献条目"{short}" 应顶格',
                        f'实际: 首行缩进{actual_pt:.0f}pt')

    # ══════════════════════════════════════════════════════════
    # 引用标注验证
    # ══════════════════════════════════════════════════════════
    def verify_citations(self):
        """验证引用标注: 上标格式、标题中无引用标注"""
        has_citation = False
        has_title_citation = False

        for para in self.paragraphs:
            text = get_text(para)
            if not text:
                continue

            is_title = match_h1(text) or match_h2(text) or match_h3(text)

            # 检查标题中是否有引用标注
            if is_title and re.search(r'\[\d+[\d,\-]*\]', text):
                has_title_citation = True
                short = text[:25]
                self.verify('引用', False,
                            f'标题"{short}"中不应有引用标注')

            # 检查引用标注上标格式
            for run in para.runs:
                if re.search(r'\[\d+[\d,\-]*\]', run.text):
                    has_citation = True
                    if not run.font.superscript:
                        # 排除"由文献[8]可知"这类非上标引用
                        if not re.search(r'由文献|见文献|文献\[', para.text):
                            self.verify('引用', False,
                                        f'引用标注"{run.text[:20]}"应为上标格式')

        if not has_citation:
            # 没有引用标注也算通过（可能文档中没有引用）
            pass

        if not has_title_citation:
            self.verify('引用', True, '标题中无引用标注')

    # ══════════════════════════════════════════════════════════
    # 结构完整性验证
    # ══════════════════════════════════════════════════════════
    def verify_structure(self):
        """验证文档结构: 目录、摘要、正文、参考文献、致谢"""
        all_text = '\n'.join(get_text(p) for p in self.paragraphs)

        # 目录
        has_toc = '目录' in all_text or '目 录' in all_text
        if has_toc:
            self.verify('结构', True, '包含目录')
        else:
            self.verify('结构', False, '应包含目录')

        # 摘要
        has_abstract = '摘 要' in all_text or '摘要' in all_text
        if has_abstract:
            self.verify('结构', True, '包含摘要')
        else:
            self.verify('结构', False, '应包含摘要')

        # 正文 (有章标题)
        has_body = bool(re.search(r'第[一二三四五六七八九十\d]+章', all_text))
        if has_body:
            self.verify('结构', True, '包含正文(章标题)')
        else:
            self.verify('结构', False, '应包含正文(章标题)')

        # 参考文献
        has_refs = '参考文献' in all_text or '参 考 文 献' in all_text
        if has_refs:
            self.verify('结构', True, '包含参考文献')
        else:
            self.verify('结构', False, '应包含参考文献')

        # 致谢
        has_thanks = '致谢' in all_text or '致 谢' in all_text
        if has_thanks:
            self.verify('结构', True, '包含致谢')
        else:
            self.verify('结构', False, '应包含致谢')

    # ══════════════════════════════════════════════════════════
    # 运行全部验证
    # ══════════════════════════════════════════════════════════
    def run_all(self):
        """运行所有验证项并输出结果"""
        # 标题头
        print(f"\n{'='*60}")
        print(f"  湖南科技大学毕业论文格式验证")
        print(f"  文件: {os.path.basename(self.filepath)}")
        print(f"{'='*60}")

        # 执行所有验证
        self.verify_page_setup()
        self.verify_header()
        self.verify_page_number()
        self.verify_font_and_paragraph()
        self.verify_citations()
        self.verify_structure()

        # 输出每项结果
        for category, ok, label, detail in self.results:
            if ok:
                print(f"  [{category}] ✓ {label}")
            else:
                if detail:
                    print(f"  [{category}] ✗ {label} ({detail})")
                else:
                    print(f"  [{category}] ✗ {label}")

        # 汇总
        print(f"{'─'*60}")
        print(f"  通过: {self.passed}/{self.total} 项")
        print(f"  失败: {self.failed}/{self.total} 项")
        print(f"{'─'*60}")

        if self.failed == 0:
            print(f"  结果: 通过")
        else:
            print(f"  结果: 未通过")

        return self.failed == 0


# ─── 入口 ───
def main():
    if len(sys.argv) < 2:
        print("用法: python verify_thesis.py <docx文件路径>")
        print("示例: python verify_thesis.py 论文.docx")
        sys.exit(1)

    filepath = sys.argv[1]
    if not os.path.exists(filepath):
        print(f"错误: 文件不存在 - {filepath}")
        sys.exit(1)

    verifier = ThesisVerifier(filepath)
    ok = verifier.run_all()
    sys.exit(0 if ok else 1)


if __name__ == '__main__':
    main()