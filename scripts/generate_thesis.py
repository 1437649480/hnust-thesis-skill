# -*- coding: utf-8 -*-
"""
湖南科技大学本科毕业论文生成脚本
根据规范生成符合格式要求的示例docx文件（基于附件6-11示例内容）
用法: python generate_thesis.py [输出路径，默认: 格式示例.docx]
"""

import sys, os, copy, re
from docx import Document
from docx.shared import Pt, Cm, Inches, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─── 常量 ───
SPEC = {
    'page': {'width': Cm(21), 'height': Cm(29.7), 'top': Cm(3), 'bottom': Cm(2.5),
             'left': Cm(2.5), 'right': Cm(2.5), 'header': Cm(2), 'footer': Cm(1.5)},
}

def set_page(section):
    """设置页面"""
    section.page_width = SPEC['page']['width']
    section.page_height = SPEC['page']['height']
    section.top_margin = SPEC['page']['top']
    section.bottom_margin = SPEC['page']['bottom']
    section.left_margin = SPEC['page']['left']
    section.right_margin = SPEC['page']['right']
    section.header_distance = SPEC['page']['header']
    section.footer_distance = SPEC['page']['footer']

def set_run_font(run, font_name, size, bold=False, east_asian=None):
    """设置run字体，同时设置中西文字体"""
    run.font.name = font_name
    run.font.size = size
    run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), east_asian or font_name)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)

def add_paragraph(doc, text, font_name='宋体', size=Pt(12), bold=False, align=None,
                  space_before=0, space_after=0, line_spacing=1.25, first_line_indent=None,
                  font_name_east=None):
    """添加格式化段落"""
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before) if isinstance(space_before, (int, float)) else space_before
    pf.space_after = Pt(space_after) if isinstance(space_after, (int, float)) else space_after
    pf.line_spacing = line_spacing
    if first_line_indent:
        pf.first_line_indent = first_line_indent

    run = p.add_run(text)
    set_run_font(run, font_name, size, bold, font_name_east)
    return p

def add_title(doc, text, level='h1'):
    """添加标题"""
    if level == 'h1':
        return add_paragraph(doc, text, '宋体', Pt(18), True, WD_ALIGN_PARAGRAPH.CENTER, 6, 6, 1.25)
    elif level == 'h2':
        return add_paragraph(doc, text, '宋体', Pt(14), True, WD_ALIGN_PARAGRAPH.LEFT, 6, 6, 1.25)
    elif level == 'h3':
        return add_paragraph(doc, text, '宋体', Pt(12), True, WD_ALIGN_PARAGRAPH.LEFT, 6, 6, 1.25)
    elif level == 'h4':
        return add_paragraph(doc, text, '宋体', Pt(12), True, WD_ALIGN_PARAGRAPH.LEFT, 0, 0, 1.25, Pt(24))

def add_body(doc, text):
    """添加正文段落"""
    return add_paragraph(doc, text, '宋体', Pt(12), False, WD_ALIGN_PARAGRAPH.LEFT, 0, 0, 1.25, Pt(24))

def add_ref(doc, text):
    """添加参考文献条目（中文宋体五号，英文Times New Roman五号）"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.25
    # 使用中英文混合字体：中文宋体，英文Times New Roman
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = 'Times New Roman'
    run.bold = False
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    return p

def add_page_break(doc):
    """添加分页符"""
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._element.append(br)

def setup_header(section):
    """设置页眉"""
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.clear()
    run = p.add_run('湖南科技大学本科生毕业设计（论文）')
    set_run_font(run, '宋体', Pt(10.5), False)

def setup_footer(section):
    """设置页脚（PAGE域）"""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.clear()

    r1 = p.add_run('-')
    set_run_font(r1, 'Times New Roman', Pt(10.5))

    r2 = p.add_run()
    set_run_font(r2, 'Times New Roman', Pt(10.5))
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    r2._element.append(fld_begin)
    instr = OxmlElement('w:instrText')
    instr.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    instr.text = ' PAGE '
    r2._element.append(instr)
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    r2._element.append(fld_end)

    r3 = p.add_run('-')
    set_run_font(r3, 'Times New Roman', Pt(10.5))

def add_table_example(doc):
    """添加示例表格（表2.1 1978-2007年我国外汇储备情况）"""
    # 表名
    add_paragraph(doc, '表2.1 1978—2007年我国外汇储备情况', '宋体', Pt(12), True, WD_ALIGN_PARAGRAPH.CENTER, 6, 0)
    # 单位
    add_paragraph(doc, '单位：10亿美元', '宋体', Pt(10.5), False, WD_ALIGN_PARAGRAPH.RIGHT, 0, 0)

    data = [
        ['年末', '外汇储备', '年末', '外汇储备', '年末', '外汇储备'],
        ['1978', '0.167', '1988', '3.372', '1998', '144.959'],
        ['1979', '0.840', '1989', '5.550', '1999', '154.675'],
        ['1980', '-1.296', '1990', '11.093', '2000', '165.574'],
        ['1981', '2.708', '1991', '21.712', '2001', '212.165'],
        ['1982', '6.986', '1992', '19.443', '2002', '286.407'],
        ['1983', '8.901', '1993', '21.199', '2003', '403.251'],
        ['1984', '8.220', '1994', '51.620', '2004', '609.932'],
        ['1985', '2.644', '1995', '73.597', '2005', '818.872'],
        ['1986', '2.072', '1996', '105.049', '2006', '1066.344'],
        ['1987', '2.923', '1997', '139.890', '2007', '1528.249'],
    ]

    table = doc.add_table(rows=len(data), cols=6)
    table.style = 'Table Grid'

    # 设置表头字体
    for i, row_data in enumerate(data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(cell_text)
            if i == 0:
                set_run_font(run, '宋体', Pt(10.5), True)
            else:
                set_run_font(run, 'Times New Roman', Pt(10.5))

    # 资料来源
    add_paragraph(doc, '资料来源：中国国家外汇管理局网站www.safe.gov.cn', '宋体', Pt(10.5), False, WD_ALIGN_PARAGRAPH.LEFT, 0, 6)

def generate(filepath):
    """生成完整示例文档"""
    doc = Document()

    # ── 页面设置 ──
    section = doc.sections[0]
    set_page(section)
    setup_header(section)
    setup_footer(section)

    # ── 扉页（附件1） ──
    add_paragraph(doc, '湖 南 科 技 大 学', '宋体', Pt(22), True, WD_ALIGN_PARAGRAPH.CENTER, 0, 12)
    add_paragraph(doc, '毕 业 设 计（ 论 文 ）', '宋体', Pt(22), True, WD_ALIGN_PARAGRAPH.CENTER, 0, 24)

    # 扉页表格
    cover_data = [['题目', ''], ['作者', ''], ['学院', ''], ['专业', ''], ['学号', ''], ['指导教师', '']]
    table = doc.add_table(rows=len(cover_data), cols=2)
    for i, (label, val) in enumerate(cover_data):
        row = table.rows[i]
        row.cells[0].text = ''
        p = row.cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(label)
        set_run_font(run, '宋体', Pt(22), True)
        row.cells[1].text = ''

    add_paragraph(doc, '', '宋体', Pt(12), False, WD_ALIGN_PARAGRAPH.CENTER, 12, 0)
    add_paragraph(doc, '二〇〇   年   月   日', '宋体', Pt(22), True, WD_ALIGN_PARAGRAPH.CENTER, 6, 0)

    add_page_break(doc)

    # ── 任务书（附件2） ──
    add_paragraph(doc, '湖 南 科 技 大 学', '宋体', Pt(22), True, WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, '毕业设计（论文）任务书', '宋体', Pt(22), True, WD_ALIGN_PARAGRAPH.CENTER, 0, 12)
    add_body(doc, '院  系（教研室）')
    add_body(doc, '系（教研室）主任: （签名）              年  月  日')
    add_body(doc, '学生姓名:              学号:              专业:')
    add_body(doc, '1 设计（论文）题目及专题：')
    add_body(doc, '2 学生设计（论文）时间：自    年    月    日开始至    年    月    日止')
    add_body(doc, '3 设计（论文）所用资源和参考资料：')
    add_body(doc, '4 设计（论文）应完成的主要内容：')
    add_body(doc, '5 提交设计（论文）形式（设计说明与图纸或论文等）及要求：')
    add_body(doc, '6 发题时间：    年    月    日')
    add_body(doc, '指导教师：        （签名）')
    add_body(doc, '学    生：        （签名）')

    add_page_break(doc)

    # ── 指导人评语（附件3） ──
    add_paragraph(doc, '湖 南 科 技 大 学', '宋体', Pt(22), True, WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, '毕业设计（论文）指导人评语', '宋体', Pt(22), True, WD_ALIGN_PARAGRAPH.CENTER, 0, 12)
    add_body(doc, '主要对学生毕业设计（论文）的工作态度，研究内容与方法，工作量，文献应用，创新性，实用性，科学性，文本（图纸）规范程度，存在的不足等进行综合评价。')
    add_paragraph(doc, '', '宋体', Pt(12), False, WD_ALIGN_PARAGRAPH.LEFT, 24, 0)
    add_body(doc, '指导人：        （签名）')
    add_body(doc, '年  月  日')
    add_body(doc, '指导人评定成绩：')

    add_page_break(doc)

    # ── 评阅人评语（附件4） ──
    add_paragraph(doc, '湖 南 科 技 大 学', '宋体', Pt(22), True, WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, '毕业设计（论文）评阅人评语', '宋体', Pt(22), True, WD_ALIGN_PARAGRAPH.CENTER, 0, 12)
    add_body(doc, '主要对学生毕业设计（论文）的文本格式、图纸规范程度，工作量，研究内容与方法，实用性与科学性，结论和存在的不足等进行综合评价。')
    add_paragraph(doc, '', '宋体', Pt(12), False, WD_ALIGN_PARAGRAPH.LEFT, 24, 0)
    add_body(doc, '评阅人：        （签名）')
    add_body(doc, '年  月  日')
    add_body(doc, '评阅人评定成绩：')

    add_page_break(doc)

    # ── 答辩记录（附件5） ──
    add_paragraph(doc, '湖 南 科 技 大 学', '宋体', Pt(22), True, WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, '毕业设计（论文）答辩记录', '宋体', Pt(22), True, WD_ALIGN_PARAGRAPH.CENTER, 0, 12)
    add_body(doc, '日期：')
    add_body(doc, '学生：        学号：        班级：')
    add_body(doc, '题目：')
    add_body(doc, '提交毕业设计（论文）答辩委员会下列材料：')
    add_body(doc, '1 设计（论文）说明书    共  页')
    add_body(doc, '2 设计（论文）图  纸    共  页')
    add_body(doc, '3 指导人、评阅人评语   共  页')
    add_body(doc, '毕业设计（论文）答辩委员会评语：')
    add_body(doc, '主要对学生毕业设计（论文）的研究思路，设计（论文）质量，文本图纸规范程度和对设计（论文）的介绍，回答问题情况等进行综合评价。')
    add_paragraph(doc, '', '宋体', Pt(12), False, WD_ALIGN_PARAGRAPH.LEFT, 24, 0)
    add_body(doc, '答辩委员会主任：        （签名）')
    add_body(doc, '委员：        （签名）')
    add_body(doc, '              （签名）')
    add_body(doc, '              （签名）')
    add_body(doc, '              （签名）')
    add_body(doc, '答辩成绩：')
    add_body(doc, '总评成绩：')

    add_page_break(doc)

    # ── 中文摘要（附件6） ──
    add_paragraph(doc, '摘  要', '宋体', Pt(18), True, WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, '', '宋体', Pt(12), False, WD_ALIGN_PARAGRAPH.LEFT, 0, 0, 1.0)
    add_body(doc, '近三十年来，我国外汇储备数量有了很大的增长，从1996年突破1000亿美元到2007年底突破1.5万亿美元。如此高速的增长和巨大的规模，对我国经济的发展而言是把双刃剑。我国作为一个发展中国家，势必需要一定量的外汇储备来确保我国有能力对外支付、干预外汇市场以及提升国家信誉。但是过量的外汇储备规模又会产生管理性问题。本文首先介绍了我国外汇储备的发展状况；分析了外汇储备快速增长的原因；阐述了高额外汇储备的负面效应；揭示了我国外汇储备管理中存在的问题；接着就五个国家的外汇储备管理体系作了简要分析；最后对我国外汇储备的管理提出了几点建议。')
    add_paragraph(doc, '', '宋体', Pt(12), False, WD_ALIGN_PARAGRAPH.LEFT, 0, 0, 1.0)

    # 关键词行
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run('关键词：')
    set_run_font(r1, '黑体', Pt(12), False, '黑体')
    r2 = p.add_run('外汇储备；国际比较；启示')
    set_run_font(r2, '宋体', Pt(12))

    add_page_break(doc)

    # ── 英文摘要（附件6） ──
    add_paragraph(doc, 'ABSTRACT', 'Times New Roman', Pt(18), True, WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, '', 'Times New Roman', Pt(12), False, WD_ALIGN_PARAGRAPH.LEFT, 0, 0, 1.0)
    add_paragraph(doc, 'The number of our foreign exchange reserves has experienced quick growth over the past thirty years, from breaking 100 billion U.S. dollar in 1996 to breaking 1.5 trillion U.S. dollar in 2007. The so high-speed increase and gigantic scale, are a double-edged sword for our country development of the economy. As a developing country, China definitely needs a certain amount of foreign exchange to ensure external payments, interfere in foreign exchange market and promote the country credit. But excessive exchange cover scale may produce problems in management. At first, this paper introduces the development of our country foreign exchange reserves, analyses the reasons for quick growth of foreign exchange reserves, illustrates the negative effects of excessive foreign exchange reserves, reveals the problems in foreign exchange reserve; Then analyses and comprises the management systems of foreign exchange reserves in five counties; At last, makes recommendations for our country\'s management of foreign exchange reserves.',
                  'Times New Roman', Pt(12), False, WD_ALIGN_PARAGRAPH.LEFT, 0, 0, 1.25, Pt(24))
    add_paragraph(doc, '', 'Times New Roman', Pt(12), False, WD_ALIGN_PARAGRAPH.LEFT, 0, 0, 1.0)

    # Keywords
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    r1 = p.add_run('Keywords: ')
    set_run_font(r1, 'Times New Roman', Pt(12), True)
    r2 = p.add_run('foreign exchange reserves; international comparison; meanings')
    set_run_font(r2, 'Times New Roman', Pt(12), False)

    # ── 目录（附件7） ──
    # 需要在目录前插入分节符，以便目录页码为罗马数字
    # 由于 Document 默认只有一个节，我们需要在目录前也插入分节符
    # 简化处理：封面等前置部分和后置部分各一个节
    add_page_break(doc)

    add_paragraph(doc, '目  录', '黑体', Pt(16), False, WD_ALIGN_PARAGRAPH.CENTER, 0, 12)

    toc_items = [
        ('第一章 前言', '1'),
        ('第二章 我国外汇储备的发展阶段', '1'),
        ('2.1 规模较小阶段', '1'),
        ('2.2 较快增长阶段', '1'),
        ('2.3 缓慢增长阶段', '2'),
        ('2.4 大幅度增长阶段', '2'),
        ('第三章 结论', '15'),
        ('参考文献', '16'),
        ('致谢', '17'),
        ('附录A', '18'),
    ]
    for item, page in toc_items:
        is_chapter = item.startswith('第')
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.tab_stops.add_tab_stop(Cm(14.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT, leader=1)
        r1 = p.add_run(item)
        set_run_font(r1, '宋体', Pt(14) if is_chapter and '第' in item else Pt(12), is_chapter and '第' in item)
        r2 = p.add_run('\t')
        r3 = p.add_run(page)
        set_run_font(r3, 'Times New Roman', Pt(12))

    add_page_break(doc)

    # ── 正文（附件8） ──
    add_title(doc, '第一章 前  言', 'h1')
    add_body(doc, '外汇储备是指国际储备中的各种能充当储备货币的资产，它是货币行政当局以银行存款、财政部库存、长短期政府证券等形式所保有的，在国际收支逆差时可以使用的债权。我国的外汇储备主要有美元、欧元、日元、英镑等。')
    add_body(doc, '一国的外汇储备，必须具备四个基本特征：第一，为国家直接持有；第二，是国际通行的可自由兑换货币；第三，储备资产必须具有流动性的性质；第四，其主要作用是用于平衡国际收支和稳定汇率。为了分析外汇储备的来源结构，还可以将外汇储备划分为债权性储备和债务性储备，前者由商品出口、劳务出口等创汇形成，在国际收支平衡表中反映为经常项目顺差；后者由国外借款、外商直接投资及国际游资构成，在国际收支平衡表中反映为资本和金融项目顺差。两者的比例，反映了一个国家外汇储备的质量。')
    add_body(doc, '外汇储备与货币当局的黄金储备、在国际货币基金组织的头寸、特别提款权及其他债权一起，构成一国或地区的国际储备。外汇储备是国际储备中规模最大、增长最快、地位最重要的资产，占国际储备资产总额的绝大比重。')
    add_body(doc, '一定的外汇储备是一国进行经济调节、实现内外平衡的重要手段。当国际收支出现逆差时，动用外汇储备可以促进国际收支的平衡；当国内宏观经济不平衡，出现总需求大于总供给时，可以动用外汇组织进口，从而调节总供给与总需求的关系，促进宏观经济的平衡。同时当汇率出现波动时，可以利用外汇储备干预汇率，使之趋于稳定。')

    add_page_break(doc)

    add_title(doc, '第二章 我国外汇储备的发展阶段', 'h1')
    add_body(doc, '改革开放以来，我国外汇储备的增长大体经历了以下四个阶段：')

    add_title(doc, '2.1 规模较小阶段（1978年－1993年）', 'h2')
    add_body(doc, '1978年，我国外汇储备只有16亿美元。改革开放以后，通过努力增加出口，控制进口，我国外汇储备逐渐增加，1983年达到89亿美元。当时我国认为储备过多，导致之后几年外汇储备急剧减少，到1986年下降至21亿美元，此后，逐渐恢复并一直维持在一二百亿美元的水平。')

    add_title(doc, '2.2 较快增长阶段（1994年－1997年）', 'h2')
    add_body(doc, '1994年，我国对外汇管理体制进行了重大改革，实施了汇率并轨、取消外汇留成、实行银行结售汇制、成立银行间外汇交易市场等举措，国家外汇储备获得了较快的增长。至1997年底，我国外汇储备余额由1993年的211.99亿美元增长到1398.9亿美元，增加了5.6倍，国家外汇储备进入了较为宽松的时期。')

    add_title(doc, '2.3 缓慢增长阶段（1998年－2000年）', 'h2')
    add_body(doc, '1997年下半年，亚洲金融危机爆发。从1998年起，我国外汇储备增量明显减缓。1998—2000年，国家外汇储备年增长额仅为50.97亿美元、97.15亿美元、108.99亿美元，只相当于1997年外汇储备增加额的14.62％、27.87％、31.27％。尽管如此，至2000年末，外汇储备还是增加到1655.74亿美元，居世界前列。')

    add_title(doc, '2.4 大幅度增长阶段（2001年至今）', 'h2')
    add_body(doc, '从2001年起，我国外汇储备进入了大幅度增长阶段，且增长速度惊人。2001年－2007年，国家外汇储备年增长额分别为465.91亿美元、742.42亿美元、1168.44亿美元、2066.81亿美元、2089亿美元、10663.44亿美元和15282.49亿美元。至2008年3月末，我国外汇储备更高达16822亿美元（表1）。')

    # 添加表格示例
    add_table_example(doc)

    add_page_break(doc)

    add_title(doc, '第五章 结  论', 'h1')
    add_body(doc, '外汇储备管理是国家在健全的储备管理体系下，持有适度的储备量并进行有效的管理，以实现外汇储备各项职能的一系列工作，它与国家的经济开放程度密切相关。随着世界经济一体化进程的步伐加快，国际间资金流动更加频繁，在这种形势下，各国对外汇储备的需求不仅是维持国际收支平衡的需要，更主要是抵消国际资金冲击的需要，所以说目前各国的外汇储备从某种意义上成了干预储备或安全储备。所以无论积极外汇储备管理模式具体操作如何，它都是以宏观经济长期稳定发展和人民福利提高为最终目标的。')

    add_page_break(doc)

    # ── 参考文献（附件9） ──
    add_paragraph(doc, '参 考 文 献', '宋体', Pt(14), True, WD_ALIGN_PARAGRAPH.CENTER, 0, 6)

    refs = [
        '[1] 袁庆龙,候文义.Ni-P合金镀层组织形貌及显微硬度研究[J].太原理工大学学报,2001,32(1):51-53.',
        '[2] 刘国钧,郑如斯.中国书的故事[M].北京:中国青年出版社,1979:115.',
        '[3] 孙品一.高校学报编辑工作现代化特征[C].中国高等学校自然科学学报研究会.科技编辑学论文集.北京:北京师范大学出版社,1998:10-22.',
        '[4] 张和生.地质力学系统理论[D].太原:太原理工大学,1998.',
        '[5] 冯西桥.核反应堆压力容器的LBB分析[R].北京:清华大学核能技术设计研究院,1997.',
        '[6] 姜锡洲.一种温热外敷药制备方案[P].中国专利:881056078,1983-08-12.',
        '[7] GB/T 16159-1996.汉语拼音正词法基本规则[S].北京:中国标准出版社,1996.',
        '[8] 谢希德.创造学习的思路[N].人民日报,1998-12-25(10).',
        '[9] 姚伯元.中国学术期刊标准化数据库系统工程[EB/OL].http://www.cajcd.cn/pub/wml.txt/9808.html,1998-08-16/1998-10-04.',
    ]
    for ref in refs:
        add_ref(doc, ref)

    add_page_break(doc)

    # ── 致谢（附件10） ──
    add_paragraph(doc, '致  谢', '宋体', Pt(14), True, WD_ALIGN_PARAGRAPH.CENTER, 0, 6)
    add_body(doc, '从论文选题到搜集资料，从提纲的完成到正文的反复修改，我经历了喜悦、聒噪、痛苦和彷徨，在写作论文的过程中，心情是如此复杂。如今，伴随着这篇毕业论文的最终成稿，复杂的心情烟消云散，自己甚至还有一点成就感。')
    add_body(doc, '我要感谢我的导师XXX老师和XXX老师。他们为人随和热情，治学严谨细心。从选题、定题、撰写提纲，到论文的反复修改、润色直至定稿，两位老师始终认真负责地给予我深刻而细致地指导。正是有了老师们的无私帮助与热忱鼓励，我的毕业论文才得以顺利完成。')
    add_body(doc, '我还要感谢我的班主任XXX老师以及在大学四年中给我们授课的所有老师们，是他们让我学到了很多很多知识，让我看到了世界的精彩，让我学会了做人做事。')
    add_body(doc, '最后感谢四年里陪伴我的同学、朋友们，有了他们我的人生才丰富，有了他们我在奋斗的路上才不孤独，谢谢他们。')

    add_page_break(doc)

    # ── 附录（附件11插图示例 + 附录A） ──
    # 插图示例
    add_paragraph(doc, '图3.1 催化剂的XRD图谱', '宋体', Pt(12), True, WD_ALIGN_PARAGRAPH.CENTER, 0, 6)
    # 插入一个占位图像框
    add_paragraph(doc, '[此处插入XRD图谱图片]', '宋体', Pt(10.5), False, WD_ALIGN_PARAGRAPH.CENTER, 0, 0)
    add_paragraph(doc, '', '宋体', Pt(12), False, WD_ALIGN_PARAGRAPH.LEFT, 0, 0, 1.0)

    add_page_break(doc)

    # 附录A
    add_title(doc, '附录A：主要源程序清单', 'h1')
    add_body(doc, '本附录列出了毕业设计中使用的主要程序代码片段。')
    add_paragraph(doc, '', '宋体', Pt(12), False, WD_ALIGN_PARAGRAPH.LEFT, 0, 0, 1.0)
    add_body(doc, '（附录编排格式与正文相同，此处为示例占位。）')

    # ── 保存 ──
    doc.save(filepath)
    print(f"  ✓ 示例文档已保存到: {filepath}")
    return filepath

def main():
    outpath = sys.argv[1] if len(sys.argv) > 1 else '格式示例.docx'
    # 确保输出到 examples 目录
    if not os.path.dirname(outpath):
        outpath = os.path.join(os.path.dirname(__file__), '..', 'examples', outpath)
    generate(outpath)

if __name__ == '__main__':
    main()