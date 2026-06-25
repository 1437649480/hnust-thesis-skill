# -*- coding: utf-8 -*-
"""
湖南科技大学本科毕业论文页码分节修复脚本
在目录与正文之间插入分节符，设置前置部分罗马数字页码、正文阿拉伯数字页码
页码格式为 "-X-"（含 PAGE 域代码），Times New Roman 五号居中
用法: python fix_pagenum.py <docx文件路径>
"""

import sys, os, re, copy
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─── 工具函数 ───

def detect_toc(doc):
    """检测文档中是否有"目录"段落（匹配"目 录"或"目录"）"""
    for para in doc.paragraphs:
        text = para.text.strip()
        if text == '目录' or text == '目 录':
            return True
    return False


def find_chapter1(doc):
    """找到第一章段落，返回段落对象。匹配"第X章"（X为中文数字或阿拉伯数字）"""
    pattern = re.compile(r'^第[一二三四五六七八九十\d]+章')
    for para in doc.paragraphs:
        if pattern.match(para.text.strip()):
            return para
    return None


def insert_section_break(doc, chapter1_para):
    """
    在第一章之前插入分节符（XML 级别操作）。
    返回 True 表示成功插入，False 表示失败或已存在分节符。
    """
    body = doc.element.body

    # 获取 body 中所有子元素（按文档顺序）
    children = list(body)

    # 找到第一章段落在 body 子元素中的位置
    chapter1_element = chapter1_para._element
    chapter1_idx = None
    for i, child in enumerate(children):
        if child is chapter1_element:
            chapter1_idx = i
            break

    if chapter1_idx is None:
        print("  ⚠ 警告: 无法定位第一章段落")
        return False

    # 找到第一章之前的最后一个段落（p 元素）
    prev_para = None
    for i in range(chapter1_idx - 1, -1, -1):
        if children[i].tag == qn('w:p'):
            prev_para = children[i]
            break

    if prev_para is None:
        print("  ⚠ 警告: 第一章之前没有段落，无法插入分节符")
        return False

    # 检查该段落是否已包含 sectPr（避免重复分节）
    pPr_existing = prev_para.find(qn('w:pPr'))
    if pPr_existing is not None and pPr_existing.find(qn('w:sectPr')) is not None:
        print("  ✓ 第一章之前已有分节符，跳过")
        return False

    # 获取 body 末尾的 w:sectPr 元素
    original_sectPr = body.find(qn('w:sectPr'))
    if original_sectPr is None:
        print("  ⚠ 警告: 未找到 body 的 sectPr 元素")
        return False

    # 克隆 body 末尾的 w:sectPr 元素
    cloned_sectPr = copy.deepcopy(original_sectPr)

    # 修改克隆的 sectPr: 设置页码为小写罗马数字
    pgNumType = cloned_sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        cloned_sectPr.insert(0, pgNumType)
    pgNumType.set(qn('w:fmt'), 'lowerRoman')
    # 罗马数字不需要 start 属性，移除可能存在的 start
    if qn('w:start') in pgNumType.attrib:
        del pgNumType.attrib[qn('w:start')]

    # 确保克隆的 sectPr 中断开页眉页脚链接（新建节默认独立）
    # 移除 titlePg 如果存在（不需要首页不同）
    titlePg = cloned_sectPr.find(qn('w:titlePg'))
    if titlePg is not None:
        cloned_sectPr.remove(titlePg)

    # 将克隆的 sectPr 插入到前一个段落的 w:pPr 中
    pPr = prev_para.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        prev_para.insert(0, pPr)
    pPr.append(cloned_sectPr)

    # 修改原 sectPr: 设置页码为阿拉伯数字，起始值为 1
    pgNumType_orig = original_sectPr.find(qn('w:pgNumType'))
    if pgNumType_orig is None:
        pgNumType_orig = OxmlElement('w:pgNumType')
        original_sectPr.insert(0, pgNumType_orig)
    pgNumType_orig.set(qn('w:fmt'), 'decimal')
    pgNumType_orig.set(qn('w:start'), '1')

    return True


def _make_run_properties():
    """创建页码 run 的字体属性：Times New Roman, 五号(10.5pt)"""
    rPr = OxmlElement('w:rPr')
    # 字体设置
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    rPr.append(rFonts)
    # 字号：10.5pt = 21 half-points
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '21')
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), '21')
    rPr.append(szCs)
    return rPr


def _make_text_run(text):
    """创建包含文本的 w:r 元素"""
    r = OxmlElement('w:r')
    r.append(_make_run_properties())
    t = OxmlElement('w:t')
    t.set(qn('xml:space'), 'preserve')
    t.text = text
    r.append(t)
    return r


def _make_fldChar_run(fldCharType):
    """创建 fldChar 的 w:r 元素（begin/end）"""
    r = OxmlElement('w:r')
    r.append(_make_run_properties())
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), fldCharType)
    r.append(fldChar)
    return r


def _make_instrText_run():
    """创建 instrText ' PAGE ' 的 w:r 元素"""
    r = OxmlElement('w:r')
    r.append(_make_run_properties())
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    r.append(instrText)
    return r


def setup_footer_for_section(section):
    """
    为单个节设置页脚为 "-X-" 格式（含 PAGE 域代码）。
    断开 is_linked_to_previous，清空原页脚，重新构建。
    """
    footer = section.footer

    # 断开"链接到前一节"
    footer.is_linked_to_previous = False

    # 重新获取 footer（断开链接后可能创建了新的 footer 部件）
    footer = section.footer

    # 删除所有现有段落
    footer_element = footer._element
    for p in footer_element.findall(qn('w:p')):
        footer_element.remove(p)
    for tbl in footer_element.findall(qn('w:tbl')):
        footer_element.remove(tbl)

    # 创建新段落
    new_para = OxmlElement('w:p')

    # 段落属性：居中
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    new_para.append(pPr)

    # 添加 run: "-"
    new_para.append(_make_text_run('-'))

    # 添加 PAGE 域代码：fldChar begin
    new_para.append(_make_fldChar_run('begin'))

    # 添加 PAGE 域代码：instrText " PAGE "
    new_para.append(_make_instrText_run())

    # 添加 PAGE 域代码：fldChar end
    new_para.append(_make_fldChar_run('end'))

    # 添加 run: "-"
    new_para.append(_make_text_run('-'))

    # 将新段落添加到页脚
    footer_element.append(new_para)


def ensure_pgNumType(section, fmt, start=None):
    """确保节的页码格式正确"""
    sectPr = section._sectPr
    pgNumType = sectPr.find(qn('w:pgNumType'))
    if pgNumType is None:
        pgNumType = OxmlElement('w:pgNumType')
        sectPr.insert(0, pgNumType)
    pgNumType.set(qn('w:fmt'), fmt)
    if start is not None:
        pgNumType.set(qn('w:start'), str(start))
    elif qn('w:start') in pgNumType.attrib:
        del pgNumType.attrib[qn('w:start')]


# ─── 主流程 ───

def main():
    if len(sys.argv) < 2:
        print("用法: python fix_pagenum.py <docx文件路径>")
        print("示例: python fix_pagenum.py 论文.docx")
        sys.exit(1)

    filepath = sys.argv[1]
    if not os.path.exists(filepath):
        print(f"错误: 文件不存在 - {filepath}")
        sys.exit(1)

    # 生成输出路径（在原文件名后加 _页码修复）
    base, ext = os.path.splitext(filepath)
    output_path = f"{base}_页码修复{ext}"

    print(f"\n{'='*60}")
    print(f"  湖南科技大学毕业论文页码分节修复")
    print(f"  文件: {os.path.basename(filepath)}")
    print(f"{'='*60}\n")

    # ─── 第一步：打开文档，检测并插入分节符 ───
    print("【第一步】检测目录并插入分节符")
    doc = Document(filepath)

    has_toc = detect_toc(doc)
    print(f"  目录检测: {'有' if has_toc else '无'}")

    section_inserted = False
    if has_toc:
        chapter1 = find_chapter1(doc)
        if chapter1:
            title = chapter1.text.strip()
            print(f"  第一章: {title[:40]}{'...' if len(title) > 40 else ''}")
            section_inserted = insert_section_break(doc, chapter1)
            if section_inserted:
                print("  ✓ 分节符已插入（目录 → 正文）")
            else:
                print("  分节符未插入（可能已存在或无需分节）")
        else:
            print("  ⚠ 未找到第一章，无法分节")

    # 保存文档（带分节信息）
    doc.save(output_path)
    print(f"  已保存: {os.path.basename(output_path)}\n")

    # ─── 第二步：重新打开文档，设置页脚 ───
    print("【第二步】设置页脚格式")
    doc2 = Document(output_path)
    sections = doc2.sections

    print(f"  文档共有 {len(sections)} 个节")

    for i, section in enumerate(sections):
        # 确定页码格式：节1（如有分节）为罗马数字，其余为阿拉伯数字
        is_roman = (section_inserted and i == 0)
        fmt_name = '罗马数字 (lowerRoman)' if is_roman else '阿拉伯数字 (decimal)'

        print(f"  节{i + 1}: {fmt_name}")

        # 确保页码类型正确
        if is_roman:
            ensure_pgNumType(section, 'lowerRoman')
        else:
            ensure_pgNumType(section, 'decimal', start=1)

        # 设置页脚为 "-X-" 格式
        setup_footer_for_section(section)
        print(f"    页脚已设置为 \"-X-\" 格式")

    # 保存最终文档
    doc2.save(output_path)

    # ─── 输出结果 ───
    print(f"\n{'='*60}")
    print(f"  处理完成!")
    print(f"  输出文件: {os.path.basename(output_path)}")
    if section_inserted:
        print(f"  节1（目录等前置部分）: 罗马数字页码 -i-, -ii-, -iii-...")
        print(f"  节2（正文部分）: 阿拉伯数字页码 -1-, -2-, -3-...")
    else:
        print(f"  未分节，全文使用阿拉伯数字页码 -1-, -2-, -3-...")
    print(f"{'='*60}\n")


if __name__ == '__main__':
    main()