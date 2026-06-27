# -*- coding: utf-8 -*-
"""
湖南科技大学本科毕业论文格式检查脚本
覆盖规范第3.2~3.7节、第4.1~4.9节所有可自动检查项
用法: python check_thesis.py <docx文件路径>
"""

import sys, re, os
from docx import Document
from docx.shared import Pt, Cm, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# ─── 常量定义 ───
SPEC = {
    'page': {'width': Cm(21), 'height': Cm(29.7), 'top': Cm(3), 'bottom': Cm(2.5),
             'left': Cm(2.5), 'right': Cm(2.5), 'header': Cm(2), 'footer': Cm(1.5)},
    'header': {'text': '湖南科技大学本科生毕业设计（论文）', 'font': '宋体', 'size': Pt(10.5), 'align': WD_ALIGN_PARAGRAPH.CENTER},
    'page_num': {'font': 'Times New Roman', 'size': Pt(10.5), 'align': WD_ALIGN_PARAGRAPH.CENTER},
    'h1': {'font': '宋体', 'size': Pt(18), 'bold': True, 'align': WD_ALIGN_PARAGRAPH.CENTER, 'before': Pt(6), 'after': Pt(6)},
    'h2': {'font': '宋体', 'size': Pt(14), 'bold': True, 'align': None, 'before': Pt(6), 'after': Pt(6)},
    'h3': {'font': '宋体', 'size': Pt(12), 'bold': True, 'align': None, 'before': Pt(6), 'after': Pt(6)},
    'h4': {'font': '宋体', 'size': Pt(12), 'bold': True, 'align': None, 'before': Pt(0), 'after': Pt(0)},
    'body': {'font': '宋体', 'size': Pt(12), 'bold': False, 'align': None, 'before': Pt(0), 'after': Pt(0),
             'line_spacing': 1.25, 'first_line_indent': Pt(24)},
    'abstract_title': {'font': '宋体', 'size': Pt(18), 'bold': True, 'align': WD_ALIGN_PARAGRAPH.CENTER},
    'abstract_body': {'font': '宋体', 'size': Pt(12), 'bold': False},
    'abstract_kw_label': {'font': '黑体', 'size': Pt(12), 'bold': False},
    'abstract_kw': {'font': '宋体', 'size': Pt(12), 'bold': False},
    'en_abstract_title': {'font': 'Times New Roman', 'size': Pt(18), 'bold': True, 'align': WD_ALIGN_PARAGRAPH.CENTER},
    'en_abstract_body': {'font': 'Times New Roman', 'size': Pt(12), 'bold': False},
    'en_kw_label': {'font': 'Times New Roman', 'size': Pt(12), 'bold': True},
    'en_kw': {'font': 'Times New Roman', 'size': Pt(12), 'bold': False},
    'toc_title': {'font': '黑体', 'size': Pt(16), 'bold': False},
    'ref_title': {'font': '宋体', 'size': Pt(14), 'bold': True, 'align': WD_ALIGN_PARAGRAPH.CENTER},
    'ref_entry': {'font': '宋体', 'size': Pt(10.5), 'bold': False},
    'ref_entry_en': {'font': 'Times New Roman', 'size': Pt(10.5), 'bold': False},
    'cover': {'font': '宋体', 'size': Pt(22), 'bold': True},
    'thanks_title': {'font': '宋体', 'size': Pt(14), 'bold': True, 'align': WD_ALIGN_PARAGRAPH.CENTER},
    'appendix_title': {'font': '宋体', 'size': Pt(18), 'bold': True, 'align': WD_ALIGN_PARAGRAPH.CENTER},
    'table_cell': {'font_cn': '宋体', 'font_en': 'Times New Roman', 'size': Pt(10.5)},
}

TOLERANCE = 50000  # EMU 容差

# ─── 工具函数 ───
def get_text(para):
    """获取段落纯文本"""
    return para.text.strip()

def get_run_info(para):
    """获取段落中第一个run的字体信息"""
    for run in para.runs:
        if run.text.strip():
            return run
    return para.runs[0] if para.runs else None

def match_h1(text):
    """匹配章标题: 第X章"""
    return bool(re.match(r'^第[一二三四五六七八九十\d]+章\s', text))

def match_h2(text):
    """匹配节标题: X.Y（不匹配X.Y.Z）"""
    return bool(re.match(r'^[\d]+\.[\d]+(?!\.[\d])\s', text))

def match_h3(text):
    """匹配条标题: X.Y.Z 或 X.Y.Z.W"""
    return bool(re.match(r'^[\d]+\.[\d]+\.[\d]+(\.[\d]+)?\s', text))

def match_h4(text):
    """匹配款/项标题: (1) (2) 或 ① 等"""
    return bool(re.match(r'^[（(][\d一二三四五六七八九十]+[）)]', text))

def match_ref(text):
    """匹配参考文献条目: [1] 开头"""
    return bool(re.match(r'^\[\d+\]', text))

def match_appendix(text):
    """匹配附录标题: 附录A 或 附录A："""
    return bool(re.match(r'^附录[A-Z]', text))

def match_thanks_title(text):
    """匹配致谢标题"""
    return text in ('致 谢', '致谢') or bool(re.match(r'^致\s+谢$', text))

def emu_to_cm(v):
    return v / 360000

def pt_to_cm(pt):
    return pt / 72 * 2.54

def check_val(actual, expected, label, tolerance=TOLERANCE):
    """检查数值是否在容差范围内"""
    if isinstance(expected, (int, float)) and isinstance(actual, (int, float)):
        return abs(actual - expected) <= tolerance
    return actual == expected

# ─── 检查器类 ───
class ThesisChecker:
    def __init__(self, filepath):
        self.filepath = filepath
        self.doc = Document(filepath)
        self.errors = []
        self.warnings = []
        self.passed = []
        self.sections = self.doc.sections
        self.paragraphs = self.doc.paragraphs
        # 预扫描：识别文档区域
        self.toc_start = None
        self.toc_end = None
        self.cover_end = None
        self._scan_regions()

    def check(self, condition, category, detail):
        """记录检查结果"""
        if condition:
            self.passed.append(f"[{category}] {detail}")
        else:
            self.errors.append(f"[{category}] {detail}")

    def _scan_regions(self):
        """预扫描文档区域：识别目录、扉页等边界"""
        all_text = [get_text(p) for p in self.paragraphs]
        # 找到目录起始（"目 录" 或 "目录"）
        for i, t in enumerate(all_text):
            if t in ('目 录', '目录') or re.match(r'^目\s+录$', t):
                self.toc_start = i
                break
        # 找到目录结束（目录后第一个不含页码的"第X章"）
        # 目录中的条目通常包含制表符和页码，如 "第一章 前言\t1"
        # 真正的章标题不含页码
        if self.toc_start is not None:
            for i in range(self.toc_start + 1, len(all_text)):
                text = all_text[i]
                if match_h1(text) and '\t' not in text:
                    # 不含制表符的章标题才是真正的正文标题
                    self.toc_end = i
                    break
        # 找到扉页结束（第一个"摘  要"或"ABSTRACT"）
        for i, t in enumerate(all_text):
            if t in ('摘  要', '摘 要', '摘要') or re.match(r'^摘\s+要$', t) or t.strip().upper() == 'ABSTRACT':
                self.cover_end = i
                break

    def _is_in_toc(self, para_idx):
        """判断段落是否在目录区域"""
        if self.toc_start is None or self.toc_end is None:
            return False
        return self.toc_start <= para_idx < self.toc_end

    def _is_in_cover(self, para_idx):
        """判断段落是否在扉页区域"""
        if self.cover_end is None:
            return False
        return para_idx < self.cover_end

    def check_warn(self, condition, category, detail):
        """记录警告（无法自动修复的问题）"""
        if condition:
            self.passed.append(f"[{category}] {detail}")
        else:
            self.warnings.append(f"[{category}] {detail}")

    # ── 4.1.1 页面设置 ──
    def check_page_setup(self):
        sec = self.sections[0]
        self.check(abs(sec.page_width - SPEC['page']['width']) <= TOLERANCE,
                   '页面', f'A4宽度 210mm (实际: {emu_to_cm(sec.page_width):.1f}mm)')
        self.check(abs(sec.page_height - SPEC['page']['height']) <= TOLERANCE,
                   '页面', f'A4高度 297mm (实际: {emu_to_cm(sec.page_height):.1f}mm)')
        self.check(abs(sec.top_margin - SPEC['page']['top']) <= TOLERANCE,
                   '页面', f'上页边距 30mm (实际: {emu_to_cm(sec.top_margin):.1f}mm)')
        self.check(abs(sec.bottom_margin - SPEC['page']['bottom']) <= TOLERANCE,
                   '页面', f'下页边距 25mm (实际: {emu_to_cm(sec.bottom_margin):.1f}mm)')
        self.check(abs(sec.left_margin - SPEC['page']['left']) <= TOLERANCE,
                   '页面', f'左页边距 25mm (实际: {emu_to_cm(sec.left_margin):.1f}mm)')
        self.check(abs(sec.right_margin - SPEC['page']['right']) <= TOLERANCE,
                   '页面', f'右页边距 25mm (实际: {emu_to_cm(sec.right_margin):.1f}mm)')
        self.check(abs(sec.header_distance - SPEC['page']['header']) <= TOLERANCE,
                   '页面', f'页眉距边界 20mm (实际: {emu_to_cm(sec.header_distance):.1f}mm)')
        self.check(abs(sec.footer_distance - SPEC['page']['footer']) <= TOLERANCE,
                   '页面', f'页脚距边界 15mm (实际: {emu_to_cm(sec.footer_distance):.1f}mm)')

    # ── 4.1.2 页眉 ──
    def check_header(self):
        for si, sec in enumerate(self.sections):
            header = sec.header
            if header.is_linked_to_previous:
                self.check(False, f'页眉(节{si+1})', '页眉未断开"链接到前一节"')
                continue
            paras = header.paragraphs
            if not paras or not paras[0].text.strip():
                self.check(False, f'页眉(节{si+1})', '页眉为空')
                continue
            p = paras[0]
            text = p.text.strip()
            self.check(SPEC['header']['text'] in text,
                       f'页眉(节{si+1})', f'内容应为"{SPEC["header"]["text"]}" (实际: "{text}")')
            self.check(p.alignment == SPEC['header']['align'],
                       f'页眉(节{si+1})', '页眉应居中')
            run = get_run_info(p)
            if run:
                self.check(run.font.size and abs(run.font.size - SPEC['header']['size']) <= TOLERANCE,
                           f'页眉(节{si+1})', f'字号应为五号/10.5pt (实际: {run.font.size})')
                self.check(run.font.name == SPEC['header']['font'],
                           f'页眉(节{si+1})', f'字体应为宋体 (实际: {run.font.name})')

    # ── 4.1.3 页码 ──
    def check_page_number(self):
        for si, sec in enumerate(self.sections):
            footer = sec.footer
            sectPr = sec._sectPr
            pgnt = sectPr.find(qn('w:pgNumType'))
            if pgnt is not None:
                fmt = pgnt.get(qn('w:fmt'))
                if si == 0 and len(self.sections) > 1:
                    self.check(fmt == 'lowerRoman',
                               f'页码(节{si+1})', f'前置部分应为罗马数字 (实际: {fmt})')
                else:
                    self.check(fmt == 'decimal',
                               f'页码(节{si+1})', f'正文部分应为阿拉伯数字 (实际: {fmt})')
            # 检查页脚格式
            if not footer.is_linked_to_previous:
                paras = footer.paragraphs
                if paras:
                    p = paras[0]
                    self.check(p.alignment == SPEC['page_num']['align'],
                               f'页码(节{si+1})', '页码应居中')
                    run = get_run_info(p)
                    if run:
                        self.check(run.font.size and abs(run.font.size - SPEC['page_num']['size']) <= TOLERANCE,
                                   f'页码(节{si+1})', f'页码字号应为五号/10.5pt (实际: {run.font.size})')

    # ── 4.2 字体字号段落 ──
    def check_font_paragraph(self):
        for i, para in enumerate(self.paragraphs):
            text = get_text(para)
            if not text:
                continue

            # 跳过目录区域
            if self._is_in_toc(i):
                continue
            # 跳过扉页区域
            if self._is_in_cover(i):
                continue

            style = para.style.name if para.style else ''

            # 章标题
            if match_h1(text):
                self.check_font(para, 'h1', f'章标题: {text[:30]}')
                self.check(para.paragraph_format.alignment == SPEC['h1']['align'],
                           '格式', f'章标题"{text[:20]}"应居中')
                self.check(para.paragraph_format.space_before is None or
                           abs(para.paragraph_format.space_before - SPEC['h1']['before']) <= TOLERANCE,
                           '格式', f'章标题"{text[:20]}"段前应0.5行(6pt)')
                self.check(para.paragraph_format.space_after is None or
                           abs(para.paragraph_format.space_after - SPEC['h1']['after']) <= TOLERANCE,
                           '格式', f'章标题"{text[:20]}"段后应0.5行(6pt)')
                # 标题长度检查
                title_only = re.sub(r'^第[一二三四五六七八九十\d]+章\s*', '', text)
                self.check_warn(len(title_only) <= 15,
                                '标题', f'章标题"{text[:20]}"应≤15字 (实际: {len(title_only)}字)')
                self.check_warn(not re.search(r'[，。；：？！、]', title_only),
                                '标题', f'章标题"{text[:20]}"不应使用标点符号')
                continue

            # 节标题
            if match_h2(text):
                self.check_font(para, 'h2', f'节标题: {text[:30]}')
                continue

            # 条标题
            if match_h3(text):
                self.check_font(para, 'h3', f'条标题: {text[:30]}')
                continue

            # 款/项标题
            if match_h4(text):
                self.check_font(para, 'h4', f'款/项标题: {text[:30]}')
                self.check(para.paragraph_format.first_line_indent is not None,
                           '格式', f'款/项标题"{text[:20]}"应有首行缩进')
                continue

            # 参考文献标题
            if text in ('参考文献', '参 考 文 献'):
                self.check_font(para, 'ref_title', '参考文献标题')
                self.check(para.paragraph_format.alignment == SPEC['ref_title']['align'],
                           '格式', '参考文献标题应居中')
                continue

            # 参考文献条目
            if match_ref(text):
                self.check_font(para, 'ref_entry', f'参考文献: {text[:40]}')
                self.check(para.paragraph_format.first_line_indent is None or
                           para.paragraph_format.first_line_indent == 0,
                           '格式', f'参考文献"{text[:30]}"应顶格')
                # 检查条目中英文部分字体
                self._check_ref_entry_mixed_font(para, text)
                continue

            # 致谢标题
            if match_thanks_title(text):
                self.check_font(para, 'thanks_title', '致谢标题')
                self.check(para.paragraph_format.alignment == SPEC['thanks_title']['align'],
                           '格式', '致谢标题应居中')
                continue

            # 附录标题
            if match_appendix(text):
                self.check(para.paragraph_format.alignment == SPEC['appendix_title']['align'],
                           '格式', f'附录标题"{text[:20]}"应居中')
                # 附录标题字号应与章标题相同
                run = get_run_info(para)
                if run:
                    self.check(run.font.size and abs(run.font.size - SPEC['appendix_title']['size']) <= TOLERANCE,
                               '字号', f'附录标题"{text[:20]}"字号应为小二号/18pt')
                    self.check(run.bold, '加粗', f'附录标题"{text[:20]}"应加粗')
                continue

            # 摘要标题
            if text in ('摘 要', '摘要'):
                self.check_font(para, 'abstract_title', '中文摘要标题')
                self.check(para.paragraph_format.alignment == SPEC['abstract_title']['align'],
                           '格式', '摘要标题应居中')
                continue

            # 英文摘要标题
            if text.strip().upper() == 'ABSTRACT':
                self.check_font(para, 'en_abstract_title', '英文摘要标题')
                self.check(para.paragraph_format.alignment == SPEC['en_abstract_title']['align'],
                           '格式', 'ABSTRACT应居中')
                continue

            # 关键词
            if text.startswith('关键词') or text.startswith('关键词：'):
                self.check_kw_format(para)
                continue
            if text.startswith('Keywords') or text.startswith('KeyWords'):
                self.check_en_kw_format(para)
                continue

            # 正文内容
            self.check_body_format(para, text)

    def check_font(self, para, level, label):
        spec = SPEC[level]
        run = get_run_info(para)
        if not run:
            return
        # 参考文献条目特殊处理：西文TNR，中文宋体
        if level == 'ref_entry':
            # 检查西文字体
            rPr = run._element.find(qn('w:rPr'))
            ascii_font = None
            ea_font = None
            if rPr is not None:
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is not None:
                    ascii_font = rFonts.get(qn('w:ascii'))
                    ea_font = rFonts.get(qn('w:eastAsia'))
            self.check(ascii_font == 'Times New Roman' or run.font.name == 'Times New Roman',
                       '字体', f'{label} 西文字体应为Times New Roman (实际: {ascii_font or run.font.name})')
            self.check(ea_font == '宋体' or run.font.name == '宋体',
                       '字体', f'{label} 中文字体应为宋体 (实际: {ea_font or run.font.name})')
            self.check(run.font.size and abs(run.font.size - spec['size']) <= TOLERANCE,
                       '字号', f'{label} 字号应为{spec["size"]/12700:.0f}pt')
            return
        self.check(run.font.name and run.font.name == spec['font'],
                   '字体', f'{label} 字体应为{spec["font"]} (实际: {run.font.name})')
        self.check(run.font.size and abs(run.font.size - spec['size']) <= TOLERANCE,
                   '字号', f'{label} 字号应为{spec["size"]/12700:.0f}pt (实际: {run.font.size/12700 if run.font.size else "None"}pt)')
        if spec.get('bold'):
            self.check(run.bold,
                       '加粗', f'{label} 应为加粗')

    def check_body_format(self, para, text):
        """检查正文格式：宋体小四号，段前段后不空行，1.25倍行距，首行缩进2字符"""
        # 跳过表名、图题、致谢标题等非正文段落
        if re.match(r'^(表|图)[\d]+\.[\d]+', text):
            return  # 表名/图题，不检查正文格式
        if text.startswith('资料来源') or text.startswith('单位：'):
            return  # 表格说明文字
        if match_thanks_title(text):
            return  # 致谢标题
        
        # 字体字号检查
        run = get_run_info(para)
        if run:
            if run.font.name:
                self.check(run.font.name == SPEC['body']['font'],
                           '字体', f'正文"{text[:20]}"字体应为宋体 (实际: {run.font.name})')
            if run.font.size:
                self.check(abs(run.font.size - SPEC['body']['size']) <= TOLERANCE,
                           '字号', f'正文"{text[:20]}"字号应为小四号/12pt (实际: {run.font.size/12700 if run.font.size else "None"}pt)')
        
        # 行距
        pf = para.paragraph_format
        line_spacing = pf.line_spacing
        if line_spacing is not None:
            self.check(abs(line_spacing - SPEC['body']['line_spacing']) <= 0.05,
                       '行距', f'正文"{text[:20]}"行距应为1.25倍 (实际: {line_spacing})')
        # 首行缩进
        if pf.first_line_indent is not None:
            self.check(abs(pf.first_line_indent - SPEC['body']['first_line_indent']) <= TOLERANCE,
                       '缩进', f'正文"{text[:20]}"首行缩进应为2字符/24pt (实际: {pf.first_line_indent/12700 if pf.first_line_indent else "None"}pt)')
        # 段前段后
        if pf.space_before is not None and pf.space_before > 0:
            self.check(False, '段间距', f'正文"{text[:20]}"段前不应有空行 (实际: {pf.space_before/12700:.0f}pt)')
        if pf.space_after is not None and pf.space_after > 0:
            self.check(False, '段间距', f'正文"{text[:20]}"段后不应有空行 (实际: {pf.space_after/12700:.0f}pt)')

    def check_kw_format(self, para):
        """检查中文关键词格式"""
        text = get_text(para)
        # 关键词标签应为黑体
        for run in para.runs:
            if '关键词' in run.text:
                self.check(run.font.name == '黑体',
                           '字体', f'关键词标签应黑体 (实际: {run.font.name})')
                self.check(run.font.size and abs(run.font.size - Pt(12)) <= TOLERANCE,
                           '字号', '关键词标签应小四号/12pt')

    def check_en_kw_format(self, para):
        """检查英文关键词格式"""
        for run in para.runs:
            if 'Keyword' in run.text or 'KEYWORD' in run.text.upper():
                self.check(run.font.name == 'Times New Roman',
                           '字体', f'Keywords标签应Times New Roman (实际: {run.font.name})')
                self.check(run.bold,
                           '加粗', 'Keywords标签应加粗')

    # ── 3.3 引用标注 ──
    def check_citations(self):
        """检查引用文献标注格式"""
        in_refs = False
        for para in self.paragraphs:
            text = get_text(para)
            if not text:
                continue
            # 进入/离开参考文献区域
            if text in ('参考文献', '参 考 文 献'):
                in_refs = True
                continue
            if match_thanks_title(text) or match_appendix(text):
                in_refs = False
                continue
            if in_refs:
                continue  # 跳过参考文献条目
            # 跳过款/项标题（以（1）、（2）等开头）
            if re.match(r'^[（(][\d一二三四五六七八九十]+[）)]', text):
                continue
            # 跳过标题
            if match_h1(text) or match_h2(text) or match_h3(text):
                # 检查标题中是否有引用标注（不应有）
                if re.search(r'\[\d+\]', text):
                    self.check(False, '引用', f'标题"{text[:20]}"中不应有引用标注')
                continue

            # 检查上标引用
            for run in para.runs:
                if re.search(r'\[\d+[\d,\-]*\]', run.text):
                    if not run.font.superscript:
                        # 可能是正文中的引用，如"由文献[8]可知"
                        if not re.search(r'由文献|见文献', para.text):
                            self.check(False, '引用', f'引用标注"{run.text[:20]}"应为上标格式')

    # ── 参考文献条目中英文混排检查 ──
    def _check_ref_entry_mixed_font(self, para, text):
        """检查参考文献条目中英文部分字体"""
        for run in para.runs:
            run_text = run.text.strip()
            if not run_text:
                continue
            # 检查是否包含英文字符
            has_english = bool(re.search(r'[a-zA-Z]', run_text))
            if has_english:
                # 英文部分应用Times New Roman
                # 检查 run.font.name 或 XML 中的 ascii 字体
                rPr = run._element.find(qn('w:rPr'))
                ascii_font = run.font.name
                if rPr is not None:
                    rFonts = rPr.find(qn('w:rFonts'))
                    if rFonts is not None:
                        ascii_font = rFonts.get(qn('w:ascii')) or ascii_font
                if ascii_font and ascii_font != 'Times New Roman':
                    self.check_warn(False, '字体',
                                    f'参考文献条目中英文部分"{run_text[:20]}"应用Times New Roman (实际: {ascii_font})')

    # ── 关键词数量检查 ──
    def check_keyword_count(self):
        """检查关键词数量是否为3~5个"""
        for para in self.paragraphs:
            text = get_text(para)
            # 中文关键词
            if text.startswith('关键词') or text.startswith('关键词：'):
                kw_text = re.sub(r'^关键词[：:]?\s*', '', text)
                keywords = [k.strip() for k in kw_text.split('；') if k.strip()]
                if not keywords:
                    keywords = [k.strip() for k in kw_text.split(';') if k.strip()]
                self.check(3 <= len(keywords) <= 5,
                           '关键词', f'中文关键词应为3~5个 (实际: {len(keywords)}个)')
            # 英文关键词
            if text.lower().startswith('keywords'):
                kw_text = re.sub(r'^keywords[：:]\s*', '', text, flags=re.IGNORECASE)
                keywords = [k.strip() for k in kw_text.split(';') if k.strip()]
                if not keywords:
                    keywords = [k.strip() for k in kw_text.split('；') if k.strip()]
                self.check(3 <= len(keywords) <= 5,
                           '关键词', f'英文关键词应为3~5个 (实际: {len(keywords)}个)')

    # ── 章编号大写数字检查 (4.3) ──
    def check_chapter_number_style(self):
        """检查章标题是否使用大写数字（如"第三章"而非"第3章"）"""
        for para in self.paragraphs:
            text = get_text(para)
            if match_h1(text):
                # 检查是否使用了阿拉伯数字
                if re.match(r'^第\d+章', text):
                    self.check_warn(False, '标题', f'章标题"{text[:20]}"应使用大写数字（如"第三章"而非"第3章"）')

    # ── 论文题目检查 (2.1) ──
    def check_paper_title(self):
        """检查论文题目：≤20字，不设副标题"""
        # 在扉页区域查找题目
        for i, para in enumerate(self.paragraphs):
            text = get_text(para)
            if not text:
                continue
            # 查找扉页表格中的题目行
            if text == '题目':
                # 题目在表格中，检查下一个单元格
                # 由于表格结构复杂，这里检查扉页文本中是否有 "题目" 关键词
                pass
        # 更简单的方法：检查文档前几页是否有不合理的超长标题
        for para in self.paragraphs[:20]:
            text = get_text(para)
            if not text:
                continue
            # 检查是否有副标题（破折号或冒号分隔的标题）
            if match_h1(text):
                title_only = re.sub(r'^第[一二三四五六七八九十\d]+章\s*', '', text)
                if '——' in title_only or '——' in title_only:
                    self.check_warn(False, '标题', f'章标题"{text[:20]}"不应设副标题')

    # ── 表内文字格式检查 (4.5) ──
    def check_table_cell_format(self):
        """检查表内文字格式：中文五号宋体，英文五号Times New Roman"""
        for ti, table in enumerate(self.doc.tables):
            # 跳过扉页表格
            if ti == 0 and len(table.rows) <= 10:
                first_cell_text = ''
                if table.rows[0].cells:
                    first_cell_text = table.rows[0].cells[0].text.strip()
                if first_cell_text in ('题目', '作者', '学院'):
                    continue
            for ri, row in enumerate(table.rows):
                for ci, cell in enumerate(row.cells):
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if not run.text.strip():
                                continue
                            run_text = run.text.strip()
                            has_chinese = bool(re.search(r'[\u4e00-\u9fff]', run_text))
                            has_english = bool(re.search(r'[a-zA-Z]', run_text))
                            # 检查字号
                            if run.font.size:
                                self.check(abs(run.font.size - SPEC['table_cell']['size']) <= TOLERANCE,
                                           '表格', f'表{ti+1}单元格"{run_text[:15]}"字号应为五号/10.5pt')
                            # 中文部分检查宋体
                            if has_chinese:
                                font_name = run.font.name
                                if font_name:
                                    self.check(font_name == SPEC['table_cell']['font_cn'],
                                               '表格', f'表{ti+1}单元格中文"{run_text[:15]}"应宋体 (实际: {font_name})')

    # ── 另起页完整检查 (4.10.1) ──
    def check_all_page_breaks(self):
        """检查所有需要另起页的部分：摘要、目录、每章、参考文献、附录、致谢"""
        all_text = [get_text(p) for p in self.paragraphs]
        for i, para in enumerate(self.paragraphs):
            text = get_text(para)
            if not text:
                continue

            # 跳过目录区域
            if self._is_in_toc(i):
                continue
            # 跳过扉页区域
            if self._is_in_cover(i):
                continue

            # 需要另起页的标题
            needs_break = (
                text in ('摘 要', '摘要') or re.match(r'^摘\s+要$', text) or
                text.strip().upper() == 'ABSTRACT' or
                text in ('目 录', '目录') or re.match(r'^目\s+录$', text) or
                match_h1(text) or
                text in ('参考文献', '参 考 文 献') or
                match_thanks_title(text) or
                match_appendix(text)
            )

            if needs_break and i > 0:
                # 检查前一个段落是否有分页符或分节符
                prev_para = self.paragraphs[i - 1]
                pPr = para._element.find(qn('w:pPr'))
                has_page_break = False
                if pPr is not None:
                    page_break = pPr.find(qn('w:pageBreakBefore'))
                    sect_pr = pPr.find(qn('w:sectPr'))
                    has_page_break = page_break is not None or sect_pr is not None

                if not has_page_break:
                    # 也检查前一段落末尾是否有分页符或分节符
                    prev_pPr = prev_para._element.find(qn('w:pPr'))
                    if prev_pPr is not None:
                        prev_sect = prev_pPr.find(qn('w:sectPr'))
                        has_page_break = prev_sect is not None
                    # 也检查前一段落中是否有 <w:br w:type="page"/>
                    if not has_page_break:
                        for run_elem in prev_para._element.findall(qn('w:r')):
                            for br in run_elem.findall(qn('w:br')):
                                if br.get(qn('w:type')) == 'page':
                                    has_page_break = True
                                    break
                            if has_page_break:
                                break

                if not has_page_break:
                    section_name = {
                        '摘 要': '中文摘要', '摘要': '中文摘要', 'ABSTRACT': '英文摘要',
                        '目 录': '目录', '目录': '目录',
                        '参考文献': '参考文献', '参 考 文 献': '参考文献',
                        '致 谢': '致谢', '致谢': '致谢',
                    }.get(text, None)
                    if section_name is None:
                        # 尝试正则匹配
                        if re.match(r'^摘\s+要$', text):
                            section_name = '中文摘要'
                        elif re.match(r'^目\s+录$', text):
                            section_name = '目录'
                        elif match_h1(text):
                            section_name = '每章'
                        elif match_appendix(text):
                            section_name = '附录'
                    self.check(False, '分页', f'"{section_name or text}"前应另起页')

    # ── 摘要内容字体检查 ──
    def check_abstract_body_format(self):
        """检查摘要正文和英文摘要正文的字体格式"""
        in_cn_abstract = False
        in_en_abstract = False
        for para in self.paragraphs:
            text = get_text(para)
            if text in ('摘 要', '摘要'):
                in_cn_abstract = True
                in_en_abstract = False
                continue
            if text.strip().upper() == 'ABSTRACT':
                in_cn_abstract = False
                in_en_abstract = True
                continue
            if text.startswith('关键词') or text.startswith('Keywords'):
                in_cn_abstract = False
                in_en_abstract = False
                continue
            if match_h1(text):
                in_cn_abstract = False
                in_en_abstract = False
                continue

            if in_cn_abstract and text:
                # 中文摘要内容
                run = get_run_info(para)
                if run:
                    self.check(run.font.size and abs(run.font.size - SPEC['abstract_body']['size']) <= TOLERANCE,
                               '摘要', f'中文摘要内容"{text[:20]}"字号应为小四号/12pt')
            if in_en_abstract and text:
                # 英文摘要内容
                run = get_run_info(para)
                if run:
                    self.check(run.font.size and abs(run.font.size - SPEC['en_abstract_body']['size']) <= TOLERANCE,
                               '摘要', f'英文摘要内容"{text[:20]}"字号应为小四号/12pt')

    # ── 4.4 公式 ──
    def check_formulas(self):
        """检查公式格式（检测MathType或OMML公式）"""
        for para in self.paragraphs:
            # 检查是否有公式对象
            has_math = False
            for run in para.runs:
                if run._element.findall(qn('w:object')) or run._element.findall(qn('m:oMath')):
                    has_math = True
                    break
            if has_math:
                self.check(para.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER,
                           '公式', '公式应居中')

    # ── 4.5 表格 ──
    def check_tables(self):
        """检查表格格式"""
        for ti, table in enumerate(self.doc.tables):
            # 跳过扉页表格（表格行数少且位于文档前部）
            if ti == 0 and len(table.rows) <= 10:
                # 检查是否是扉页表格（包含"题目"、"作者"等关键词）
                first_cell_text = ''
                if table.rows[0].cells:
                    first_cell_text = table.rows[0].cells[0].text.strip()
                if first_cell_text in ('题目', '作者', '学院'):
                    continue  # 跳过扉页表格

            # 检查表格上方是否有表名
            tbl_element = table._element
            prev = tbl_element.getprevious()
            table_name = ''
            if prev is not None and prev.tag == qn('w:p'):
                table_name = prev.findtext(qn('w:r') + '/' + qn('w:t'), default='')
            # 检查表名格式
            if table_name:
                # 跳过非表名的文本（如单位说明、资料来源等）
                if not re.match(r'^表[\d]+\.[\d]+', table_name):
                    # 再往上找一层
                    prev2 = prev.getprevious() if prev is not None else None
                    if prev2 is not None and prev2.tag == qn('w:p'):
                        table_name = prev2.findtext(qn('w:r') + '/' + qn('w:t'), default='')
                self.check_warn(bool(re.match(r'^表[\d]+\.[\d]+', table_name)),
                                '表格', f'表名"{table_name[:30]}"应以"表X.Y"开头')
                self.check_warn('，' not in table_name and '。' not in table_name,
                                '表格', f'表名"{table_name[:30]}"不应包含标点符号')

    # ── 4.6 插图 ──
    def check_images(self):
        """检查插图格式"""
        for para in self.paragraphs:
            for run in para.runs:
                drawings = run._element.findall(qn('w:drawing'))
                if drawings:
                    # 插图下方应有图题
                    next_para = para._element.getnext()
                    if next_para is not None and next_para.tag == qn('w:p'):
                        next_text = next_para.findtext(qn('w:r') + '/' + qn('w:t'), default='')
                        self.check_warn(bool(re.match(r'^图[\d]+\.[\d]+', next_text)),
                                        '插图', f'图题"{next_text[:30]}"应以"图X.Y"开头')

    # ── 4.3 每章另起页 ──
    def check_chapter_page_break(self):
        """检查每章是否另起页"""
        for i, para in enumerate(self.paragraphs):
            text = get_text(para)
            if match_h1(text):
                # 跳过目录区域
                if self._is_in_toc(i):
                    continue
                # 检查段落前是否有分页符
                pPr = para._element.find(qn('w:pPr'))
                has_page_break = False
                if pPr is not None:
                    page_break = pPr.find(qn('w:pageBreakBefore'))
                    sect_pr = pPr.find(qn('w:sectPr'))
                    has_page_break = page_break is not None or sect_pr is not None
                # 也检查前一段落中是否有分页符
                if not has_page_break and i > 0:
                    prev_para = self.paragraphs[i - 1]
                    for run_elem in prev_para._element.findall(qn('w:r')):
                        for br in run_elem.findall(qn('w:br')):
                            if br.get(qn('w:type')) == 'page':
                                has_page_break = True
                                break
                        if has_page_break:
                            break
                if not has_page_break:
                    self.check(False, '分页', f'章标题"{text[:20]}"前应有分页符')

    # ── 3.7 年份格式 ──
    def check_year_format(self):
        """检查年份是否写全"""
        for para in self.paragraphs:
            text = get_text(para)
            # 查找如"08年"这样的缩写
            if re.search(r'(?<!\d)\d{2}年', text):
                matches = re.findall(r'(?<!\d)(\d{2}年)', text)
                for m in matches:
                    self.check_warn(False, '年份', f'年份"{m}"应写全为四位数')

    # ── 内容结构 ──
    def check_structure(self):
        """检查文档结构完整性"""
        all_text = '\n'.join(get_text(p) for p in self.paragraphs)
        has_toc = '目录' in all_text or bool(re.search(r'目\s+录', all_text))
        has_abstract = '摘要' in all_text or bool(re.search(r'摘\s+要', all_text))
        has_en_abstract = 'ABSTRACT' in all_text.upper()
        has_intro = bool(re.search(r'第[一二三四五六七八九十]+章.*[前言引绪]', all_text))
        has_conclusion = bool(re.search(r'第[一二三四五六七八九十]+章.*[结论]', all_text))
        has_refs = '参考文献' in all_text or '参 考 文 献' in all_text
        has_thanks = '致谢' in all_text or bool(re.search(r'致\s+谢', all_text))

        self.check(has_abstract, '结构', '应包含中文摘要')
        self.check(has_en_abstract, '结构', '应包含英文摘要')
        self.check(has_toc, '结构', '应包含目录')
        self.check(has_intro, '结构', '应包含引言/前言/绪论')
        self.check(has_conclusion, '结构', '应包含结论')
        self.check(has_refs, '结构', '应包含参考文献')
        self.check(has_thanks, '结构', '应包含致谢')

    # ── 汇总 ──
    def run_all(self):
        print(f"\n{'='*60}")
        print(f"  湖南科技大学毕业论文格式检查")
        print(f"  文件: {self.filepath}")
        print(f"{'='*60}")

        self.check_page_setup()
        self.check_header()
        self.check_page_number()
        self.check_font_paragraph()
        self.check_citations()
        self.check_formulas()
        self.check_tables()
        self.check_table_cell_format()
        self.check_images()
        self.check_chapter_page_break()
        self.check_all_page_breaks()
        self.check_year_format()
        self.check_structure()
        self.check_keyword_count()
        self.check_chapter_number_style()
        self.check_paper_title()
        self.check_abstract_body_format()

        # 输出结果
        print(f"\n{'-'*60}")
        print(f"  [OK] 通过: {len(self.passed)} 项")
        if self.errors:
            print(f"  [!!] 错误: {len(self.errors)} 项")
            for e in self.errors:
                print(f"    {e}")
        if self.warnings:
            print(f"  [??] 警告: {len(self.warnings)} 项")
            for w in self.warnings:
                print(f"    {w}")
        print(f"{'-'*60}")

        return len(self.errors) == 0

def main():
    if len(sys.argv) < 2:
        print("用法: python check_thesis.py <docx文件路径>")
        print("示例: python check_thesis.py 论文.docx")
        sys.exit(1)

    filepath = sys.argv[1]
    if not os.path.exists(filepath):
        print(f"错误: 文件不存在 - {filepath}")
        sys.exit(1)

    checker = ThesisChecker(filepath)
    ok = checker.run_all()
    sys.exit(0 if ok else 1)

if __name__ == '__main__':
    main()